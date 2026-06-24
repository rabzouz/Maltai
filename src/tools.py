"""Outils de l'agent Maltai.

Chaque outil = nom + description + schema JSON (format OpenAI function
calling) + une fonction async run(args, ctx) -> str.

ctx : {"is_admin": bool} — le shell est reserve aux admins.
Les outils fichiers sont sandboxes dans data/workspace/.
"""
from __future__ import annotations

import ast
import asyncio
import csv
import datetime
import html
import io
import ipaddress
import json
import operator
import base64
import os
import re
import socket
import sys
import zipfile
import xml.etree.ElementTree as ET
from html.parser import HTMLParser
from pathlib import Path
from urllib.parse import quote, urljoin, urlparse
from typing import Any, Awaitable, Callable

import httpx

from core.config import BASE_DIR, DATA_DIR, settings
from core import plans
from src import context_compress

WORKSPACE = DATA_DIR / "workspace"
WORKSPACE.mkdir(exist_ok=True)

MAX_TOOL_OUTPUT = 6000  # caracteres renvoyes au modele


def _truncate(s: str) -> str:
    if len(s) <= MAX_TOOL_OUTPUT:
        return s
    return s[:MAX_TOOL_OUTPUT] + f"\n…[tronque, {len(s)} caracteres au total]"


def _user_workspace(ctx: dict | None = None) -> Path:
    """Retourne le workspace isole par utilisateur."""
    uid = (ctx or {}).get("user_id") or "shared"
    uid = "".join(c for c in str(uid) if c.isalnum() or c in "-_")[:32] or "shared"
    ws = WORKSPACE / uid
    ws.mkdir(parents=True, exist_ok=True)
    return ws

def _safe_path(rel: str, ctx: dict | None = None) -> Path:
    """Resout un chemin relatif DANS le workspace utilisateur, refuse toute evasion."""
    ws = _user_workspace(ctx)
    p = (ws / rel).resolve()
    if not str(p).startswith(str(ws.resolve())):
        raise ValueError("Chemin hors du workspace refuse")
    return p


# --- Calculatrice (eval AST sans danger) -----------------------------------

_OPS = {
    ast.Add: operator.add, ast.Sub: operator.sub, ast.Mult: operator.mul,
    ast.Div: operator.truediv, ast.FloorDiv: operator.floordiv,
    ast.Mod: operator.mod, ast.Pow: operator.pow, ast.USub: operator.neg,
    ast.UAdd: operator.pos,
}


def _eval_node(node: ast.AST) -> float:
    if isinstance(node, ast.Expression):
        return _eval_node(node.body)
    if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
        return node.value
    if isinstance(node, ast.BinOp) and type(node.op) in _OPS:
        return _OPS[type(node.op)](_eval_node(node.left), _eval_node(node.right))
    if isinstance(node, ast.UnaryOp) and type(node.op) in _OPS:
        return _OPS[type(node.op)](_eval_node(node.operand))
    raise ValueError("Expression non autorisee")


async def tool_calculator(args: dict, ctx: dict) -> str:
    expr = str(args.get("expression", ""))
    try:
        result = _eval_node(ast.parse(expr, mode="eval"))
        return str(result)
    except (ValueError, SyntaxError, ZeroDivisionError, OverflowError) as e:
        return f"Erreur de calcul : {e}"


# --- Fichiers (sandbox data/workspace) --------------------------------------

async def tool_list_files(args: dict, ctx: dict) -> str:
    try:
        target = _safe_path(str(args.get("path", ".")), ctx)
    except ValueError as e:
        return str(e)
    if not target.exists():
        return "Dossier inexistant"
    ws = _user_workspace(ctx)
    lines = []
    for p in sorted(target.rglob("*")):
        rel = p.relative_to(ws)
        lines.append(f"{'[D]' if p.is_dir() else '[F]'} {rel}")
        if len(lines) >= 200:
            lines.append("…")
            break
    return "\n".join(lines) or "(workspace vide)"


async def tool_read_file(args: dict, ctx: dict) -> str:
    try:
        p = _safe_path(str(args.get("path", "")), ctx)
    except ValueError as e:
        return str(e)
    if not p.is_file():
        return "Fichier introuvable"
    try:
        return _truncate(p.read_text(errors="replace"))
    except OSError as e:
        return f"Erreur lecture : {e}"


async def tool_write_file(args: dict, ctx: dict) -> str:
    try:
        p = _safe_path(str(args.get("path", "")), ctx)
    except ValueError as e:
        return str(e)
    content = str(args.get("content", ""))
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        p.write_text(content)
        return f"Ecrit : {p.relative_to(_user_workspace(ctx))} ({len(content)} caracteres)"
    except OSError as e:
        return f"Erreur ecriture : {e}"


async def tool_pdf_read(args: dict, ctx: dict) -> str:
    """Extrait le texte et les metadonnees d'un PDF du workspace utilisateur."""
    try:
        p = _safe_path(str(args.get("path", "")), ctx)
    except ValueError as e:
        return str(e)
    if not p.is_file():
        return "PDF introuvable"
    if p.suffix.lower() != ".pdf":
        return "Le fichier doit avoir l'extension .pdf"
    try:
        from pypdf import PdfReader
    except ImportError:
        return "Dependance manquante : installe `pypdf`."

    try:
        reader = PdfReader(str(p))
        if reader.is_encrypted:
            try:
                reader.decrypt(str(args.get("password", "") or ""))
            except Exception:
                return "PDF chiffre : mot de passe requis ou invalide"
        page_count = len(reader.pages)
        start_page = max(1, int(args.get("start_page") or 1))
        end_page = min(page_count, int(args.get("end_page") or page_count))
        max_chars = max(1000, min(50000, int(args.get("max_chars") or 12000)))
        if start_page > end_page:
            return f"Plage invalide : le PDF contient {page_count} page(s)"

        chunks = []
        for idx in range(start_page - 1, end_page):
            text = reader.pages[idx].extract_text() or ""
            chunks.append(f"--- Page {idx + 1} ---\n{text.strip()}")
        body = "\n\n".join(chunks).strip()
        meta = reader.metadata or {}
        info = {
            "path": str(p.relative_to(_user_workspace(ctx))).replace("\\", "/"),
            "pages": page_count,
            "extracted_pages": [start_page, end_page],
            "title": str(getattr(meta, "title", "") or ""),
            "author": str(getattr(meta, "author", "") or ""),
        }
        payload = json.dumps(info, ensure_ascii=False, indent=2)
        if len(body) > max_chars:
            body = body[:max_chars] + f"\n…[tronque, limite {max_chars} caracteres]"
        return _truncate(f"{payload}\n\nTexte extrait :\n{body or '(aucun texte extractible)'}")
    except Exception as e:
        return f"Erreur lecture PDF : {e}"


def _pdf_target_path(path: str, ctx: dict | None) -> Path:
    rel = (path or "").strip() or "exports/document.pdf"
    if not rel.lower().endswith(".pdf"):
        rel += ".pdf"
    return _safe_path(rel, ctx)


def _pdf_paragraph(text: str) -> str:
    return html.escape(text).replace("\n", "<br/>")


async def tool_pdf_create(args: dict, ctx: dict) -> str:
    """Cree un PDF simple depuis du texte ou un fichier texte du workspace."""
    title = str(args.get("title", "") or "").strip()
    content = str(args.get("content", "") or "")
    source_path = str(args.get("source_path", "") or "").strip()
    if source_path and not content:
        try:
            source = _safe_path(source_path, ctx)
        except ValueError as e:
            return str(e)
        if not source.is_file():
            return "Fichier source introuvable"
        try:
            content = source.read_text(errors="replace")
        except OSError as e:
            return f"Erreur lecture source : {e}"
    if not content.strip():
        return "Contenu vide. Fournis `content` ou `source_path`."

    try:
        target = _pdf_target_path(str(args.get("path", "") or ""), ctx)
    except ValueError as e:
        return str(e)

    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, letter
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError:
        return "Dependance manquante : installe `reportlab`."

    page_size_name = str(args.get("page_size", "A4") or "A4").strip().lower()
    page_size = letter if page_size_name in {"letter", "us-letter"} else A4
    author = str(args.get("author", "MaltaiAI") or "MaltaiAI")
    target.parent.mkdir(parents=True, exist_ok=True)

    styles = getSampleStyleSheet()
    normal = ParagraphStyle(
        "MaltaiBody",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=15,
        textColor=colors.HexColor("#1f2933"),
        spaceAfter=7,
    )
    h1 = ParagraphStyle("MaltaiH1", parent=styles["Heading1"], fontSize=20, leading=24, spaceAfter=12)
    h2 = ParagraphStyle("MaltaiH2", parent=styles["Heading2"], fontSize=15, leading=19, spaceBefore=8, spaceAfter=8)
    bullet = ParagraphStyle("MaltaiBullet", parent=normal, leftIndent=12, firstLineIndent=-8)

    story = []
    if title:
        story.append(Paragraph(_pdf_paragraph(title), styles["Title"]))
        story.append(Spacer(1, 8))

    for raw_line in content[:120000].splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            story.append(Spacer(1, 6))
            continue
        stripped = line.strip()
        if stripped.startswith("# "):
            story.append(Paragraph(_pdf_paragraph(stripped[2:].strip()), h1))
        elif stripped.startswith("## "):
            story.append(Paragraph(_pdf_paragraph(stripped[3:].strip()), h2))
        elif stripped.startswith(("- ", "* ")):
            story.append(Paragraph("- " + _pdf_paragraph(stripped[2:].strip()), bullet))
        else:
            story.append(Paragraph(_pdf_paragraph(stripped), normal))

    doc = SimpleDocTemplate(
        str(target),
        pagesize=page_size,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title or target.name,
        author=author,
    )

    def footer(canvas, doc_obj):
        canvas.setTitle(title or target.name)
        canvas.setAuthor(author)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#7b8794"))
        canvas.drawRightString(page_size[0] - 18 * mm, 10 * mm, f"Page {doc_obj.page}")

    try:
        doc.build(story, onFirstPage=footer, onLaterPages=footer)
    except Exception as e:
        return f"Erreur creation PDF : {e}"

    rel = str(target.relative_to(_user_workspace(ctx))).replace("\\", "/")
    download_url = f"/api/workspace/download?path={quote(rel)}"
    return (
        f"PDF cree : {rel}\n"
        f"Taille : {target.stat().st_size} octets\n"
        f"Lien telechargement : {download_url}\n"
        f"[Telecharger le PDF]({download_url})"
    )


async def tool_docx_read(args: dict, ctx: dict) -> str:
    try:
        p = _safe_path(str(args.get("path", "")), ctx)
    except ValueError as e:
        return str(e)
    if not p.is_file() or p.suffix.lower() != ".docx":
        return "Fichier DOCX introuvable"
    try:
        from docx import Document
    except ImportError:
        return "Dependance manquante : installe `python-docx`."
    max_chars = max(1000, min(50000, int(args.get("max_chars") or 12000)))
    try:
        doc = Document(str(p))
        parts = []
        props = doc.core_properties
        meta = {
            "path": str(p.relative_to(_user_workspace(ctx))).replace("\\", "/"),
            "title": props.title or "",
            "author": props.author or "",
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
        }
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for ti, table in enumerate(doc.tables, start=1):
            parts.append(f"\n--- Tableau {ti} ---")
            for row in table.rows[:80]:
                parts.append(" | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells))
        body = "\n".join(parts).strip()
        if len(body) > max_chars:
            body = body[:max_chars] + f"\n…[tronque, limite {max_chars} caracteres]"
        return _truncate(json.dumps(meta, ensure_ascii=False, indent=2) + "\n\nTexte extrait :\n" + (body or "(vide)"))
    except Exception as e:
        return f"Erreur lecture DOCX : {e}"


def _office_target_path(path: str, suffix: str, default_name: str, ctx: dict | None) -> Path:
    rel = (path or "").strip() or f"exports/{default_name}{suffix}"
    if not rel.lower().endswith(suffix):
        rel += suffix
    return _safe_path(rel, ctx)


async def tool_docx_create(args: dict, ctx: dict) -> str:
    title = str(args.get("title", "") or "").strip()
    content = str(args.get("content", "") or "")
    source_path = str(args.get("source_path", "") or "").strip()
    if source_path and not content:
        try:
            source = _safe_path(source_path, ctx)
        except ValueError as e:
            return str(e)
        if not source.is_file():
            return "Fichier source introuvable"
        try:
            content = source.read_text(errors="replace")
        except OSError as e:
            return f"Erreur lecture source : {e}"
    if not content.strip() and not title:
        return "Contenu vide. Fournis `content`, `source_path` ou `title`."
    try:
        from docx import Document
    except ImportError:
        return "Dependance manquante : installe `python-docx`."
    try:
        target = _office_target_path(str(args.get("path", "") or ""), ".docx", "document", ctx)
    except ValueError as e:
        return str(e)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
        doc.core_properties.title = title
    doc.core_properties.author = str(args.get("author", "MaltaiAI") or "MaltaiAI")
    for raw_line in content[:160000].splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(stripped)
    try:
        doc.save(str(target))
    except Exception as e:
        return f"Erreur creation DOCX : {e}"
    rel = str(target.relative_to(_user_workspace(ctx))).replace("\\", "/")
    return f"DOCX cree : {rel}\nTaille : {target.stat().st_size} octets\nLien telechargement : /api/workspace/download?path={quote(rel)}"


async def tool_xlsx_read(args: dict, ctx: dict) -> str:
    try:
        p = _safe_path(str(args.get("path", "")), ctx)
    except ValueError as e:
        return str(e)
    if not p.is_file() or p.suffix.lower() not in {".xlsx", ".xlsm"}:
        return "Fichier XLSX introuvable"
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Dependance manquante : installe `openpyxl`."
    max_rows = max(1, min(500, int(args.get("max_rows") or 80)))
    max_cols = max(1, min(80, int(args.get("max_cols") or 20)))
    sheet_name = str(args.get("sheet", "") or "").strip()
    wb = None
    try:
        wb = load_workbook(str(p), read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name else wb[wb.sheetnames[0]]
        rows = []
        for row in ws.iter_rows(max_row=max_rows, max_col=max_cols, values_only=True):
            rows.append(["" if v is None else v for v in row])
        meta = {
            "path": str(p.relative_to(_user_workspace(ctx))).replace("\\", "/"),
            "sheets": wb.sheetnames,
            "sheet": ws.title,
            "rows_returned": len(rows),
            "max_row": ws.max_row,
            "max_column": ws.max_column,
        }
        return _truncate(json.dumps({"meta": meta, "rows": rows}, ensure_ascii=False, indent=2, default=str))
    except KeyError:
        return f"Feuille introuvable : {sheet_name}"
    except Exception as e:
        return f"Erreur lecture XLSX : {e}"
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:
                pass


def _rows_from_args(args: dict) -> list[list[Any]]:
    rows = args.get("rows")
    if isinstance(rows, list):
        return [r if isinstance(r, list) else [r] for r in rows]
    csv_text = str(args.get("csv", "") or "")
    if csv_text.strip():
        return list(csv.reader(io.StringIO(csv_text)))
    headers = args.get("headers")
    data = args.get("data")
    out = []
    if isinstance(headers, list):
        out.append(headers)
    if isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                keys = headers if isinstance(headers, list) and headers else list(item.keys())
                if not out and keys:
                    out.append(keys)
                out.append([item.get(k, "") for k in keys])
            elif isinstance(item, list):
                out.append(item)
            else:
                out.append([item])
    return out


async def tool_xlsx_create(args: dict, ctx: dict) -> str:
    rows = _rows_from_args(args)
    if not rows:
        return "Donnees vides. Fournis `rows`, `csv` ou `data`."
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill
        from openpyxl.utils import get_column_letter
    except ImportError:
        return "Dependance manquante : installe `openpyxl`."
    try:
        target = _office_target_path(str(args.get("path", "") or ""), ".xlsx", "tableau", ctx)
    except ValueError as e:
        return str(e)
    sheet = str(args.get("sheet", "Données") or "Données")[:31]
    target.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet
    for row in rows[:5000]:
        ws.append(row[:200])
    if bool(args.get("header", True)) and ws.max_row >= 1:
        for cell in ws[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F2933")
    for col in range(1, min(ws.max_column, 30) + 1):
        width = max(10, min(42, max(len(str(ws.cell(row=r, column=col).value or "")) for r in range(1, min(ws.max_row, 60) + 1)) + 2))
        ws.column_dimensions[get_column_letter(col)].width = width
    try:
        wb.save(str(target))
    except Exception as e:
        return f"Erreur creation XLSX : {e}"
    rel = str(target.relative_to(_user_workspace(ctx))).replace("\\", "/")
    return f"XLSX cree : {rel}\nLignes : {len(rows)}\nLien telechargement : /api/workspace/download?path={quote(rel)}"


def _is_hidden_path(path: Path) -> bool:
    return any(part.startswith(".") for part in path.parts)


async def tool_zip_create(args: dict, ctx: dict) -> str:
    try:
        target = _office_target_path(str(args.get("path", "") or ""), ".zip", "archive", ctx)
    except ValueError as e:
        return str(e)
    ws = _user_workspace(ctx).resolve()
    items = args.get("files")
    if not isinstance(items, list) or not items:
        folder = str(args.get("folder", ".") or ".")
        try:
            root = _safe_path(folder, ctx)
        except ValueError as e:
            return str(e)
        if not root.exists():
            return "Dossier/fichier introuvable"
        files = [p for p in ([root] if root.is_file() else root.rglob("*")) if p.is_file()]
    else:
        files = []
        for rel in items:
            try:
                p = _safe_path(str(rel), ctx)
            except ValueError as e:
                return str(e)
            if p.is_file():
                files.append(p)
            elif p.is_dir():
                files.extend(x for x in p.rglob("*") if x.is_file())
    include_hidden = bool(args.get("include_hidden", False))
    files = [p for p in files if p.resolve() != target.resolve() and (include_hidden or not _is_hidden_path(p.relative_to(ws)))]
    max_files = max(1, min(1000, int(args.get("max_files") or 500)))
    files = files[:max_files]
    if not files:
        return "Aucun fichier a zipper"
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for p in files:
                zf.write(p, p.relative_to(ws).as_posix())
    except Exception as e:
        return f"Erreur creation ZIP : {e}"
    rel = str(target.relative_to(ws)).replace("\\", "/")
    return f"ZIP cree : {rel}\nFichiers : {len(files)}\nTaille : {target.stat().st_size} octets\nLien telechargement : /api/workspace/download?path={quote(rel)}"


async def tool_zip_extract(args: dict, ctx: dict) -> str:
    try:
        archive = _safe_path(str(args.get("path", "")), ctx)
        dest = _safe_path(str(args.get("dest", "extracted") or "extracted"), ctx)
    except ValueError as e:
        return str(e)
    if not archive.is_file() or archive.suffix.lower() != ".zip":
        return "Archive ZIP introuvable"
    overwrite = bool(args.get("overwrite", False))
    max_files = max(1, min(1000, int(args.get("max_files") or 500)))
    dest.mkdir(parents=True, exist_ok=True)
    try:
        with zipfile.ZipFile(archive, "r") as zf:
            infos = [i for i in zf.infolist() if not i.is_dir()][:max_files]
            written = 0
            for info in infos:
                target = (dest / info.filename).resolve()
                if not str(target).startswith(str(dest.resolve())):
                    return "Archive refusee : chemin dangereux"
                if target.exists() and not overwrite:
                    continue
                target.parent.mkdir(parents=True, exist_ok=True)
                with zf.open(info) as src, target.open("wb") as out:
                    out.write(src.read())
                written += 1
    except Exception as e:
        return f"Erreur extraction ZIP : {e}"
    rel = str(dest.relative_to(_user_workspace(ctx))).replace("\\", "/")
    return f"ZIP extrait dans : {rel}\nFichiers extraits : {written}"


async def tool_context_compress(args: dict, ctx: dict) -> str:
    """Compresse un gros texte/JSON/log avant de le donner au modele."""
    text = str(args.get("text", "") or "")
    path = str(args.get("path", "") or "").strip()
    if path:
        try:
            p = _safe_path(path, ctx)
        except ValueError as e:
            return str(e)
        if not p.is_file():
            return "Fichier introuvable"
        try:
            text = p.read_text(errors="replace")
        except OSError as e:
            return f"Erreur lecture : {e}"

    if not text:
        return "Texte vide. Fournis `text` ou `path`."

    mode = str(args.get("mode", "auto") or "auto")
    try:
        max_chars = int(args.get("max_chars") or context_compress.DEFAULT_MAX_CHARS)
    except (TypeError, ValueError):
        max_chars = context_compress.DEFAULT_MAX_CHARS
    result = context_compress.compress_text(text, mode=mode, max_chars=max_chars)
    return context_compress.format_compressed(result)


# --- Web ---------------------------------------------------------------------

_TAG_RE = re.compile(r"<script[\s\S]*?</script>|<style[\s\S]*?</style>|<[^>]+>")


def _strip_html(raw: str) -> str:
    text = _TAG_RE.sub(" ", raw)
    text = html.unescape(text)
    return re.sub(r"\s{2,}", " ", text).strip()


_BROWSER_STATE: dict[str, dict[str, Any]] = {}
_PW = None
_PW_BROWSER = None
_PW_SESSIONS: dict[str, dict[str, Any]] = {}


def _browser_key(ctx: dict | None) -> str:
    return str((ctx or {}).get("user_id") or "shared")


def _resolve_browser_url(url: str, ctx: dict | None) -> str:
    url = (url or "").strip()
    if url.startswith(("http://", "https://")):
        return url
    current = _BROWSER_STATE.get(_browser_key(ctx), {}).get("url", "")
    if current and url.startswith("/"):
        base = re.match(r"^https?://[^/]+", current)
        if base:
            return base.group(0) + url
    return url


def _check_browser_url_allowed(url: str, ctx: dict | None) -> str:
    target = _resolve_browser_url(url, ctx)
    if not target.startswith(("http://", "https://")):
        raise ValueError("URL invalide (http/https requis)")
    host = re.sub(r"^https?://", "", target).split("/")[0].split(":")[0]
    if not (ctx or {}).get("is_admin") and _is_private_host(host):
        raise ValueError("Refuse : hote prive/interne (reserve aux administrateurs)")
    return target


def _html_title(raw: str) -> str:
    match = re.search(r"<title[^>]*>(.*?)</title>", raw, re.I | re.S)
    return _strip_html(match.group(1)) if match else ""


def _html_links(raw: str, base_url: str, limit: int = 20) -> list[dict[str, str]]:
    links = []
    for href, label in re.findall(r'<a[^>]+href=["\']([^"\']+)["\'][^>]*>(.*?)</a>', raw, re.I | re.S):
        href = html.unescape(href.strip())
        if not href or href.startswith(("#", "javascript:", "mailto:", "tel:")):
            continue
        if href.startswith("/"):
            base = re.match(r"^https?://[^/]+", base_url)
            if base:
                href = base.group(0) + href
        elif not href.startswith(("http://", "https://")):
            href = base_url.rstrip("/") + "/" + href.lstrip("./")
        text = _strip_html(label)[:120] or href
        links.append({"text": text, "url": href})
        if len(links) >= limit:
            break
    return links


def _html_forms(raw: str, base_url: str = "", limit: int = 10) -> list[dict[str, Any]]:
    forms = []
    for index, form in enumerate(re.findall(r"<form\b[^>]*>.*?</form>", raw, re.I | re.S)[:limit]):
        method = re.search(r'\bmethod=["\']?([^"\'>\s]+)', form, re.I)
        action = re.search(r'\baction=["\']?([^"\'>\s]+)', form, re.I)
        inputs = []
        for tag in re.findall(r"<(?:input|textarea|select)\b[^>]*>", form, re.I | re.S):
            name = re.search(r'\bname=["\']([^"\']+)["\']', tag, re.I)
            if not name:
                continue
            typ = re.search(r'\btype=["\']?([^"\'>\s]+)', tag, re.I)
            value = re.search(r'\bvalue=["\']([^"\']*)["\']', tag, re.I)
            inputs.append({
                "name": html.unescape(name.group(1)),
                "type": (typ.group(1).lower() if typ else "text"),
                "value": html.unescape(value.group(1)) if value else "",
            })
        action_url = html.unescape(action.group(1)) if action else base_url
        if base_url:
            action_url = urljoin(base_url, action_url or base_url)
        forms.append({
            "index": index,
            "method": (method.group(1).upper() if method else "GET"),
            "action": action_url,
            "fields": inputs[:30],
        })
    return forms


class _ScrapeHTMLParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.nodes: list[dict[str, Any]] = []
        self.stack: list[int] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        node = {
            "tag": tag.lower(),
            "attrs": {k.lower(): (v or "") for k, v in attrs},
            "text": "",
            "children": [],
        }
        idx = len(self.nodes)
        if self.stack:
            self.nodes[self.stack[-1]]["children"].append(idx)
        self.nodes.append(node)
        if tag.lower() not in {"area", "base", "br", "col", "embed", "hr", "img", "input", "link", "meta", "param", "source", "track", "wbr"}:
            self.stack.append(idx)

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        while self.stack:
            idx = self.stack.pop()
            if self.nodes[idx]["tag"] == tag:
                break

    def handle_data(self, data: str) -> None:
        if self.stack and data.strip():
            self.nodes[self.stack[-1]]["text"] += data


def _node_text(nodes: list[dict[str, Any]], idx: int) -> str:
    node = nodes[idx]
    parts = [node.get("text", "")]
    for child in node.get("children", []):
        parts.append(_node_text(nodes, child))
    return re.sub(r"\s+", " ", html.unescape(" ".join(parts))).strip()


def _parse_selector(selector: str) -> dict[str, str]:
    selector = selector.strip()
    if not selector:
        return {}
    selector = selector.split()[0]
    parsed: dict[str, str] = {}
    attr = re.search(r"\[([A-Za-z0-9_-]+)(\^=|\$=|\*=|=)?['\"]?([^'\"]*)['\"]?\]", selector)
    if attr:
        parsed["attr"] = attr.group(1).lower()
        if attr.group(2):
            parsed["attr_op"] = attr.group(2)
            parsed["attr_value"] = attr.group(3) or ""
        selector = selector[:attr.start()] + selector[attr.end():]
    node_id = re.search(r"#([A-Za-z0-9_-]+)", selector)
    if node_id:
        parsed["id"] = node_id.group(1)
        selector = selector.replace(node_id.group(0), "")
    klass = re.search(r"\.([A-Za-z0-9_-]+)", selector)
    if klass:
        parsed["class"] = klass.group(1)
        selector = selector.replace(klass.group(0), "")
    tag = selector.strip().lower()
    if tag:
        parsed["tag"] = tag
    return parsed


def _matches_selector(node: dict[str, Any], selector: str) -> bool:
    parsed = _parse_selector(selector)
    if not parsed:
        return False
    attrs = node.get("attrs", {})
    if parsed.get("tag") and node.get("tag") != parsed["tag"]:
        return False
    if parsed.get("id") and attrs.get("id") != parsed["id"]:
        return False
    if parsed.get("class"):
        classes = set((attrs.get("class") or "").split())
        if parsed["class"] not in classes:
            return False
    if parsed.get("attr"):
        if parsed["attr"] not in attrs:
            return False
        if parsed.get("attr_value") is not None:
            actual = attrs.get(parsed["attr"], "")
            expected = parsed["attr_value"]
            op = parsed.get("attr_op", "=")
            if op == "=" and actual != expected:
                return False
            if op == "^=" and not actual.startswith(expected):
                return False
            if op == "$=" and not actual.endswith(expected):
                return False
            if op == "*=" and expected not in actual:
                return False
    return True


def _extract_by_selector(nodes: list[dict[str, Any]], selector: str, attr: str = "text", all_items: bool = False, limit: int = 20) -> Any:
    values = []
    for idx, node in enumerate(nodes):
        if not _matches_selector(node, selector):
            continue
        value = _node_text(nodes, idx) if attr == "text" else node.get("attrs", {}).get(attr.lower(), "")
        if value:
            values.append(value)
        if len(values) >= limit:
            break
    return values if all_items else (values[0] if values else "")


def _extract_tables(raw: str, limit: int = 3) -> list[dict[str, Any]]:
    tables = []
    for table_raw in re.findall(r"<table\b[^>]*>(.*?)</table>", raw, re.I | re.S)[:limit]:
        rows = []
        for row_raw in re.findall(r"<tr\b[^>]*>(.*?)</tr>", table_raw, re.I | re.S):
            cells = re.findall(r"<t[hd]\b[^>]*>(.*?)</t[hd]>", row_raw, re.I | re.S)
            row = [_strip_html(c) for c in cells]
            if row:
                rows.append(row)
        if rows:
            tables.append({"rows": rows[:30]})
    return tables


def _extract_json_ld(raw: str, limit: int = 5) -> list[Any]:
    items = []
    for block in re.findall(r'<script[^>]+type=["\']application/ld\+json["\'][^>]*>(.*?)</script>', raw, re.I | re.S)[:limit]:
        try:
            items.append(json.loads(html.unescape(block.strip())))
        except json.JSONDecodeError:
            continue
    return items


def _safe_export_filename(name: str, fmt: str) -> str:
    name = re.sub(r"[^\w.\-]+", "_", (name or "scrape").strip())[:100] or "scrape"
    ext = "." + fmt
    suffix = Path(name).suffix.lower()
    if suffix != ext:
        if suffix in {".json", ".csv", ".md", ".txt", ".html"}:
            name = name[: -len(suffix)]
        name += ext
    return name


def _scrape_to_csv(data: dict[str, Any]) -> str:
    out = io.StringIO()
    writer = csv.writer(out)
    writer.writerow(["section", "key", "value"])
    for key in ("url", "status", "content_type", "title"):
        writer.writerow(["meta", key, data.get(key, "")])
    fields = data.get("fields") or {}
    if isinstance(fields, dict):
        for key, value in fields.items():
            if isinstance(value, list):
                for item in value:
                    writer.writerow(["fields", key, item])
            else:
                writer.writerow(["fields", key, value])
    for group in ("links", "images"):
        for item in data.get(group, []) or []:
            if isinstance(item, dict):
                writer.writerow([group, item.get("text") or item.get("alt") or "", item.get("url") or item.get("src") or ""])
    headings = data.get("headings") or {}
    if isinstance(headings, dict):
        for level, values in headings.items():
            for value in values or []:
                writer.writerow(["headings", level, value])
    return out.getvalue()


def _scrape_to_markdown(data: dict[str, Any]) -> str:
    lines = [
        f"# {data.get('title') or 'Scraping web'}",
        "",
        f"- URL: {data.get('url', '')}",
        f"- Status: {data.get('status', '')}",
        f"- Content-Type: {data.get('content_type', '')}",
        "",
    ]
    fields = data.get("fields") or {}
    if isinstance(fields, dict) and fields:
        lines.extend(["## Champs extraits", ""])
        for key, value in fields.items():
            lines.append(f"### {key}")
            if isinstance(value, list):
                lines.extend([f"- {v}" for v in value])
            else:
                lines.append(str(value))
            lines.append("")
    headings = data.get("headings") or {}
    if isinstance(headings, dict) and headings:
        lines.extend(["## Titres", ""])
        for level, values in headings.items():
            if values:
                lines.append(f"### {level}")
                lines.extend([f"- {v}" for v in values])
                lines.append("")
    links = data.get("links") or []
    if links:
        lines.extend(["## Liens", ""])
        for item in links:
            lines.append(f"- [{item.get('text') or item.get('url')}]({item.get('url')})")
    return "\n".join(lines).strip() + "\n"


def _scrape_to_text(data: dict[str, Any]) -> str:
    lines = [
        str(data.get("title") or "Scraping web"),
        "=" * 72,
        f"URL: {data.get('url', '')}",
        f"Status: {data.get('status', '')}",
        "",
    ]
    fields = data.get("fields") or {}
    if isinstance(fields, dict):
        for key, value in fields.items():
            lines.append(str(key).upper())
            if isinstance(value, list):
                lines.extend([f"- {v}" for v in value])
            else:
                lines.append(str(value))
            lines.append("")
    return "\n".join(lines).strip() + "\n"


def _scrape_to_html(data: dict[str, Any]) -> str:
    title = html.escape(str(data.get("title") or "Scraping web"))
    parts = [
        "<!doctype html>",
        "<html lang=\"fr\"><head><meta charset=\"utf-8\">",
        f"<title>{title}</title>",
        "<style>body{font-family:Arial,sans-serif;max-width:980px;margin:32px auto;padding:0 18px;line-height:1.5;color:#16202a}h1,h2{color:#0b3b42}code,pre{background:#f4f6f8;padding:12px;border-radius:8px;display:block;overflow:auto}li{margin:4px 0}</style>",
        "</head><body>",
        f"<h1>{title}</h1>",
        f"<p><strong>URL:</strong> {html.escape(str(data.get('url', '')))}</p>",
        f"<p><strong>Status:</strong> {html.escape(str(data.get('status', '')))}</p>",
    ]
    fields = data.get("fields") or {}
    if isinstance(fields, dict) and fields:
        parts.append("<h2>Champs extraits</h2>")
        for key, value in fields.items():
            parts.append(f"<h3>{html.escape(str(key))}</h3>")
            if isinstance(value, list):
                parts.append("<ul>")
                parts.extend([f"<li>{html.escape(str(v))}</li>" for v in value])
                parts.append("</ul>")
            else:
                parts.append(f"<p>{html.escape(str(value))}</p>")
    parts.append("<h2>JSON complet</h2>")
    parts.append(f"<pre>{html.escape(json.dumps(data, ensure_ascii=False, indent=2))}</pre>")
    parts.append("</body></html>")
    return "\n".join(parts)


def _save_scrape_export(data: dict[str, Any], args: dict, ctx: dict) -> str:
    save_as = str(args.get("save_as", "")).strip()
    if not save_as:
        return ""
    fmt = str(args.get("format", "") or Path(save_as).suffix.lstrip(".") or "json").lower()
    if fmt not in {"json", "csv", "md", "txt", "html"}:
        fmt = "json"
    filename = _safe_export_filename(save_as, fmt)
    rel = f"exports/{filename}"
    target = _safe_path(rel, ctx)
    target.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "csv":
        content = _scrape_to_csv(data)
    elif fmt == "md":
        content = _scrape_to_markdown(data)
    elif fmt == "txt":
        content = _scrape_to_text(data)
    elif fmt == "html":
        content = _scrape_to_html(data)
    else:
        content = json.dumps(data, ensure_ascii=False, indent=2)
    target.write_text(content, encoding="utf-8")
    return rel


async def tool_web_scrape(args: dict, ctx: dict) -> str:
    url = str(args.get("url", "")).strip()
    if not url.startswith(("http://", "https://")):
        return "Erreur : URL http/https requise."
    host = re.sub(r"^https?://", "", url).split("/")[0].split(":")[0]
    if not ctx.get("is_admin") and _is_private_host(host):
        return "Refuse : hote prive/interne (reserve aux administrateurs)"
    limit = max(1, min(100, int(args.get("limit") or 30)))
    fields = args.get("fields") or {}
    include = args.get("include") or {}
    if not isinstance(fields, dict):
        fields = {}
    if not isinstance(include, dict):
        include = {}
    try:
        r = await _safe_fetch(
            "GET", url, ctx, timeout=25,
            user_agent="Mozilla/5.0 (compatible; MaltaiScraper/1.0)",
        )
        if r.status_code >= 400:
            return f"Erreur HTTP {r.status_code} sur {url}"
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur reseau : {e}"

    raw = r.text
    parser = _ScrapeHTMLParser()
    parser.feed(raw)
    nodes = parser.nodes
    result: dict[str, Any] = {
        "url": str(r.url),
        "status": r.status_code,
        "content_type": r.headers.get("content-type", ""),
        "title": _html_title(raw),
    }
    metas = {}
    for node in nodes:
        if node.get("tag") != "meta":
            continue
        attrs = node.get("attrs", {})
        key = attrs.get("name") or attrs.get("property")
        if key and attrs.get("content"):
            metas[key] = attrs["content"]
    if include.get("metadata", True):
        result["metadata"] = metas
    if include.get("headings", True):
        result["headings"] = {
            f"h{level}": [_node_text(nodes, i) for i, n in enumerate(nodes) if n.get("tag") == f"h{level}"][:limit]
            for level in range(1, 4)
        }
    if include.get("links", True):
        result["links"] = [
            {"text": _node_text(nodes, i), "url": urljoin(str(r.url), n.get("attrs", {}).get("href", ""))}
            for i, n in enumerate(nodes)
            if n.get("tag") == "a" and n.get("attrs", {}).get("href")
        ][:limit]
    if include.get("images", False):
        result["images"] = [
            {"alt": n.get("attrs", {}).get("alt", ""), "src": urljoin(str(r.url), n.get("attrs", {}).get("src", ""))}
            for n in nodes
            if n.get("tag") == "img" and n.get("attrs", {}).get("src")
        ][:limit]
    if include.get("tables", False):
        result["tables"] = _extract_tables(raw)
    if include.get("json_ld", True):
        result["json_ld"] = _extract_json_ld(raw)
    if include.get("text", False):
        result["text"] = _strip_html(raw)[:4000]

    extracted: dict[str, Any] = {}
    for name, spec in fields.items():
        if isinstance(spec, str):
            selector, attr_name, all_items = spec, "text", False
        elif isinstance(spec, dict):
            selector = str(spec.get("selector", ""))
            attr_name = str(spec.get("attr", "text"))
            all_items = bool(spec.get("all"))
        else:
            continue
        if selector:
            extracted[str(name)] = _extract_by_selector(nodes, selector, attr=attr_name, all_items=all_items, limit=limit)
    if extracted:
        result["fields"] = extracted

    saved_path = _save_scrape_export(result, args, ctx)
    if saved_path:
        result["export"] = {
            "path": saved_path,
            "download_url": f"/api/workspace/download?path={quote(saved_path)}",
            "short_url": f"/{saved_path}",
        }
    payload = json.dumps(result, ensure_ascii=False, indent=2)
    if saved_path:
        return _truncate(
            f"Fichier exporte : {saved_path}\n"
            f"Lien telechargement : /api/workspace/download?path={quote(saved_path)}\n"
            f"Lien court : /{saved_path}\n\n"
            f"{payload}"
        )
    return _truncate(payload)


async def _browser_fetch(url: str, ctx: dict | None) -> tuple[str, str, str]:
    target = _check_browser_url_allowed(url, ctx)
    try:
        r = await _safe_fetch(
            "GET", target, ctx, timeout=25,
            user_agent="Mozilla/5.0 (Maltai Browser)",
        )
        r.raise_for_status()
    except PermissionError as e:
        raise ValueError(f"Refuse : {e}") from e
    except httpx.HTTPError as e:
        raise ValueError(f"Erreur browser : {e}") from e
    return str(r.url), r.headers.get("content-type", ""), r.text


async def _get_playwright_page(ctx: dict | None):
    global _PW, _PW_BROWSER
    try:
        from playwright.async_api import async_playwright
    except ImportError as e:
        raise RuntimeError("Playwright indisponible. Redeploie l'image Docker avec les dependances browser.") from e
    if _PW is None:
        _PW = await async_playwright().start()
    if _PW_BROWSER is None or not _PW_BROWSER.is_connected():
        _PW_BROWSER = await _PW.chromium.launch(headless=True, args=["--no-sandbox"])
    key = _browser_key(ctx)
    session = _PW_SESSIONS.get(key)
    if not session:
        context = await _PW_BROWSER.new_context(
            viewport={"width": 1365, "height": 768},
            ignore_https_errors=True,
            user_agent="Mozilla/5.0 (Maltai Playwright)",
        )
        page = await context.new_page()
        session = {"context": context, "page": page}
        _PW_SESSIONS[key] = session
    return session["page"]


async def _playwright_snapshot(page, ctx: dict | None, prefix: str = "") -> str:
    title = await page.title()
    url = page.url
    try:
        text = await page.locator("body").inner_text(timeout=3000)
    except Exception:
        text = ""
    links = await page.locator("a").evaluate_all(
        """els => els.slice(0, 30).map(a => ({
            text: (a.innerText || a.textContent || '').trim().slice(0, 120),
            url: a.href || ''
        })).filter(x => x.url)"""
    )
    forms = await page.locator("form").evaluate_all(
        """forms => forms.slice(0, 10).map((form, index) => ({
            index,
            method: (form.method || 'GET').toUpperCase(),
            action: form.action || location.href,
            fields: Array.from(form.querySelectorAll('input, textarea, select')).slice(0, 30).map(el => ({
                name: el.name || '',
                type: el.type || el.tagName.toLowerCase()
            })).filter(x => x.name)
        }))"""
    )
    _BROWSER_STATE[_browser_key(ctx)] = {
        "url": url,
        "title": title,
        "content_type": "playwright",
        "text": text,
        "links": links,
        "forms": forms,
    }
    lines = []
    if prefix:
        lines.append(prefix)
    lines += [f"URL: {url}", f"Title: {title or '-'}", "", "Texte:", text[:2500]]
    if links:
        lines += ["", "Liens:"]
        lines += [f"- {(item.get('text') or item.get('url'))} -> {item.get('url')}" for item in links[:15]]
    if forms:
        lines += ["", "Formulaires:"]
        lines += [
            f"- #{form.get('index')} {form.get('method')} {form.get('action')} "
            f"fields={', '.join(field.get('name', '') for field in form.get('fields', [])) or '-'}"
            for form in forms[:6]
        ]
    return _truncate("\n".join(lines))


async def tool_browser_open(args: dict, ctx: dict) -> str:
    url = str(args.get("url", "")).strip()
    wait_until = str(args.get("wait_until", "domcontentloaded"))
    timeout = max(1000, min(int(args.get("timeout_ms", 20000)), 60000))
    try:
        target = _check_browser_url_allowed(url, ctx)
        page = await _get_playwright_page(ctx)
        await page.goto(target, wait_until=wait_until, timeout=timeout)
        return await _playwright_snapshot(page, ctx, "Page ouverte avec Playwright.")
    except Exception as e:
        return f"Erreur browser_open : {e}"


async def tool_browser_click(args: dict, ctx: dict) -> str:
    selector = str(args.get("selector", "")).strip()
    text = str(args.get("text", "")).strip()
    timeout = max(1000, min(int(args.get("timeout_ms", 10000)), 60000))
    try:
        page = await _get_playwright_page(ctx)
        if selector:
            await page.locator(selector).first.click(timeout=timeout)
        elif text:
            await page.get_by_text(text, exact=False).first.click(timeout=timeout)
        else:
            return "selector ou text requis"
        try:
            await page.wait_for_load_state("domcontentloaded", timeout=5000)
        except Exception:
            pass
        return await _playwright_snapshot(page, ctx, "Click execute.")
    except Exception as e:
        return f"Erreur browser_click : {e}"


async def tool_browser_type(args: dict, ctx: dict) -> str:
    selector = str(args.get("selector", "")).strip()
    label = str(args.get("label", "")).strip()
    text = str(args.get("text", ""))
    clear = bool(args.get("clear", True))
    timeout = max(1000, min(int(args.get("timeout_ms", 10000)), 60000))
    try:
        page = await _get_playwright_page(ctx)
        if not selector and not label:
            return "selector ou label requis"
        target = page.locator(selector).first if selector else page.get_by_label(label, exact=False).first
        if clear:
            await target.fill(text, timeout=timeout)
        else:
            await target.type(text, timeout=timeout)
        return await _playwright_snapshot(page, ctx, "Texte saisi.")
    except Exception as e:
        return f"Erreur browser_type : {e}"


async def tool_browser_screenshot(args: dict, ctx: dict) -> str:
    full_page = bool(args.get("full_page", True))
    name = str(args.get("path", "")).strip() or f"browser_{int(datetime.datetime.now().timestamp())}.png"
    if not name.lower().endswith(".png"):
        name += ".png"
    safe_name = "".join(c for c in name.replace("\\", "/").split("/")[-1] if c.isalnum() or c in "-_.") or "browser.png"
    try:
        page = await _get_playwright_page(ctx)
        out_dir = _user_workspace(ctx) / "browser_screenshots"
        out_dir.mkdir(parents=True, exist_ok=True)
        target = out_dir / safe_name
        await page.screenshot(path=str(target), full_page=full_page)
        rel = target.relative_to(_user_workspace(ctx)).as_posix()
        return f"Screenshot sauvegarde : {rel}\nURL: {page.url}"
    except Exception as e:
        return f"Erreur browser_screenshot : {e}"


async def tool_browser_navigate(args: dict, ctx: dict) -> str:
    url = str(args.get("url", "")).strip()
    try:
        final_url, ctype, raw = await _browser_fetch(url, ctx)
    except ValueError as e:
        return str(e)
    title = _html_title(raw) if "html" in ctype else ""
    text = _strip_html(raw) if "html" in ctype else raw
    state = {
        "url": final_url,
        "title": title,
        "content_type": ctype,
        "text": text,
        "links": _html_links(raw, final_url) if "html" in ctype else [],
        "forms": _html_forms(raw, final_url) if "html" in ctype else [],
    }
    _BROWSER_STATE[_browser_key(ctx)] = state
    return _truncate(
        f"URL: {final_url}\nTitle: {title or '-'}\nContent-Type: {ctype or '-'}\n\n"
        f"{text[:2500]}"
    )


async def tool_browser_snapshot(args: dict, ctx: dict) -> str:
    url = str(args.get("url", "")).strip()
    if url:
        nav = await tool_browser_navigate({"url": url}, ctx)
        if nav.startswith("Erreur") or nav.startswith("URL invalide"):
            return nav
    state = _BROWSER_STATE.get(_browser_key(ctx))
    if not state:
        return "Aucune page ouverte. Appelle browser_navigate avec une URL."
    lines = [
        f"URL: {state.get('url')}",
        f"Title: {state.get('title') or '-'}",
        f"Content-Type: {state.get('content_type') or '-'}",
        "",
        "Texte:",
        str(state.get("text") or "")[:2500],
    ]
    links = state.get("links") or []
    if links:
        lines += ["", "Liens:"]
        lines += [f"- {item['text']} -> {item['url']}" for item in links[:15]]
    forms = state.get("forms") or []
    if forms:
        lines += ["", "Formulaires:"]
        lines += [
            f"- #{f['index']} {f['method']} {f['action'] or '(page courante)'} "
            f"fields={', '.join(field['name'] for field in f['fields']) or '-'}"
            for f in forms[:6]
        ]
    return _truncate("\n".join(lines))


async def tool_browser_links(args: dict, ctx: dict) -> str:
    state = _BROWSER_STATE.get(_browser_key(ctx))
    if not state:
        return "Aucune page ouverte. Appelle browser_navigate avec une URL."
    links = state.get("links") or []
    if not links:
        return "Aucun lien trouve."
    return "\n".join(f"- {item['text']}\n  {item['url']}" for item in links[:50])


async def tool_browser_form_list(args: dict, ctx: dict) -> str:
    url = str(args.get("url", "")).strip()
    if url:
        nav = await tool_browser_navigate({"url": url}, ctx)
        if nav.startswith("Erreur") or nav.startswith("URL invalide") or nav.startswith("Refuse"):
            return nav
    state = _BROWSER_STATE.get(_browser_key(ctx))
    if not state:
        return "Aucune page ouverte. Appelle browser_navigate avec une URL."
    forms = state.get("forms") or []
    if not forms:
        return "Aucun formulaire trouve."
    lines = []
    for form in forms:
        fields = ", ".join(
            f"{field['name']}:{field['type']}" + (f"={field['value']}" if field.get("value") else "")
            for field in form.get("fields", [])
        )
        lines.append(
            f"#{form['index']} {form['method']} {form['action'] or state.get('url')}\n"
            f"  fields: {fields or '-'}"
        )
    return _truncate("\n".join(lines))


async def tool_browser_submit(args: dict, ctx: dict) -> str:
    state = _BROWSER_STATE.get(_browser_key(ctx))
    if not state:
        return "Aucune page ouverte. Appelle browser_navigate avec une URL."
    forms = state.get("forms") or []
    if not forms:
        return "Aucun formulaire sur la page actuelle."
    index = int(args.get("index", 0))
    if index < 0 or index >= len(forms):
        return f"Formulaire #{index} introuvable"
    form = forms[index]
    data = args.get("data") or {}
    if not isinstance(data, dict):
        return "data doit etre un objet JSON"
    payload = {field["name"]: field.get("value", "") for field in form.get("fields", [])}
    payload.update({str(k): str(v) for k, v in data.items()})
    method = str(args.get("method") or form.get("method") or "GET").upper()
    if method not in ("GET", "POST"):
        return "Seules les methodes GET et POST sont supportees par browser_submit"
    action = str(args.get("action") or form.get("action") or state.get("url"))
    action = urljoin(str(state.get("url")), action)
    host = re.sub(r"^https?://", "", action).split("/")[0].split(":")[0]
    if not ctx.get("is_admin") and _is_private_host(host):
        return "Refuse : hote prive/interne (reserve aux administrateurs)"
    try:
        if method == "GET":
            r = await _safe_fetch("GET", action, ctx, timeout=25, params=payload,
                                  user_agent="Mozilla/5.0 (Maltai Browser)")
        else:
            r = await _safe_fetch("POST", action, ctx, timeout=25, data=payload,
                                  user_agent="Mozilla/5.0 (Maltai Browser)")
        r.raise_for_status()
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur submit : {e}"
    ctype = r.headers.get("content-type", "")
    raw = r.text
    text = _strip_html(raw) if "html" in ctype else raw
    final_url = str(r.url)
    _BROWSER_STATE[_browser_key(ctx)] = {
        "url": final_url,
        "title": _html_title(raw) if "html" in ctype else "",
        "content_type": ctype,
        "text": text,
        "links": _html_links(raw, final_url) if "html" in ctype else [],
        "forms": _html_forms(raw, final_url) if "html" in ctype else [],
    }
    return _truncate(
        f"Submitted #{index} {method} {action}\n"
        f"Final URL: {final_url}\n"
        f"Status: {r.status_code}\n\n{text[:2500]}"
    )


async def tool_web_fetch(args: dict, ctx: dict) -> str:
    url = str(args.get("url", ""))
    if not url.startswith(("http://", "https://")):
        return "URL invalide (http/https requis)"
    try:
        r = await _safe_fetch("GET", url, ctx, timeout=20, user_agent="Maltai/0.2")
        r.raise_for_status()
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur fetch : {e}"
    ctype = r.headers.get("content-type", "")
    body = r.text
    if "html" in ctype:
        body = _strip_html(body)
    return _truncate(body)


async def tool_web_search(args: dict, ctx: dict) -> str:
    """Recherche DuckDuckGo (HTML lite, sans cle API)."""
    query = str(args.get("query", "")).strip()
    if not query:
        return "Requete vide"
    try:
        async with httpx.AsyncClient(timeout=20, follow_redirects=True) as client:
            r = await client.get(
                "https://html.duckduckgo.com/html/",
                params={"q": query},
                headers={"User-Agent": "Mozilla/5.0 (Maltai/0.2)"},
            )
            r.raise_for_status()
    except httpx.HTTPError as e:
        return f"Erreur recherche : {e}"
    results = re.findall(
        r'<a[^>]+class="result__a"[^>]+href="([^"]+)"[^>]*>(.*?)</a>',
        r.text, re.S,
    )
    snippets = re.findall(
        r'<a[^>]+class="result__snippet"[^>]*>(.*?)</a>', r.text, re.S
    )
    out = []
    for i, (href, title) in enumerate(results[:6]):
        snip = _strip_html(snippets[i]) if i < len(snippets) else ""
        out.append(f"- {_strip_html(title)}\n  {href}\n  {snip}")
    return _truncate("\n".join(out)) or "Aucun resultat"


# --- Shell (admin uniquement) ------------------------------------------------

async def tool_shell(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil shell reserve aux administrateurs"
    command = str(args.get("command", ""))
    try:
        proc = await asyncio.create_subprocess_shell(
            command,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
            cwd=str(WORKSPACE),
        )
        out, _ = await asyncio.wait_for(proc.communicate(), timeout=30)
        text = out.decode(errors="replace")
        return _truncate(f"[exit {proc.returncode}]\n{text}")
    except asyncio.TimeoutError:
        proc.kill()
        return "Timeout (30s) — commande interrompue"
    except OSError as e:
        return f"Erreur shell : {e}"


# --- Git read-only (admin uniquement) ----------------------------------------

def _safe_git_ref(value: str, default: str = "HEAD") -> str:
    ref = (value or default).strip()[:100] or default
    if ref.startswith("-") or not re.fullmatch(r"[A-Za-z0-9._/@:+-]+", ref):
        raise ValueError("Reference git invalide")
    return ref


def _safe_git_pathspec(value: str) -> str | None:
    raw = (value or "").strip().replace("\\", "/")
    if not raw:
        return None
    if raw.startswith("/") or raw.startswith("../") or "/../" in raw or raw == "..":
        raise ValueError("Chemin git invalide")
    return raw


async def _run_git(args: list[str], timeout: int = 12) -> str:
    try:
        proc = await asyncio.create_subprocess_exec(
            "git", *args,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
            cwd=str(BASE_DIR),
        )
        out, _ = await asyncio.wait_for(proc.communicate(), timeout=timeout)
        text = out.decode(errors="replace").strip()
        return _truncate(f"[exit {proc.returncode}]\n{text}" if text else f"[exit {proc.returncode}]")
    except asyncio.TimeoutError:
        proc.kill()
        return "Timeout git — commande interrompue"
    except OSError as e:
        return f"Erreur git : {e}"


def _git_repo_available() -> bool:
    return (BASE_DIR / ".git").exists()


def _build_metadata() -> str:
    commit = (
        os.getenv("MALTAI_GIT_COMMIT")
        or os.getenv("SOURCE_COMMIT")
        or os.getenv("GIT_COMMIT")
        or os.getenv("COOLIFY_GIT_COMMIT")
        or os.getenv("COMMIT_SHA")
        or ""
    ).strip()
    branch = (
        os.getenv("MALTAI_GIT_BRANCH")
        or os.getenv("SOURCE_BRANCH")
        or os.getenv("GIT_BRANCH")
        or os.getenv("COOLIFY_GIT_BRANCH")
        or ""
    ).strip()
    lines = [
        f"Application: {settings.APP_NAME} {settings.APP_VERSION}",
        "Git repository: absent dans ce conteneur Docker",
    ]
    if branch:
        lines.append(f"Branch: {branch}")
    if commit:
        lines.append(f"Commit: {commit}")
    if not branch and not commit:
        lines.append("Commit/branch: non fournis par l'environnement de deploiement")
    lines.append("")
    lines.append("Note: git_status complet est disponible sur une installation lancee depuis un clone Git.")
    lines.append("Pour Docker/Coolify, definissez MALTAI_GIT_COMMIT et MALTAI_GIT_BRANCH si vous voulez afficher ces infos.")
    return "\n".join(lines)


async def tool_git_status(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil git reserve aux administrateurs"
    if not _git_repo_available():
        return _build_metadata()
    return await _run_git(["status", "--short", "--branch"])


async def tool_git_branch(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil git reserve aux administrateurs"
    if not _git_repo_available():
        return _build_metadata()
    branch = await _run_git(["branch", "--show-current"])
    commit = await _run_git(["rev-parse", "--short", "HEAD"])
    remote = await _run_git(["remote", "-v"])
    return _truncate(f"Branche:\n{branch}\n\nCommit:\n{commit}\n\nRemote:\n{remote}")


async def tool_git_log(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil git reserve aux administrateurs"
    if not _git_repo_available():
        return _build_metadata() + "\n\nHistorique git indisponible car .git n'est pas copie dans l'image Docker."
    limit = max(1, min(30, int(args.get("limit") or 10)))
    return await _run_git(["log", "--oneline", "--decorate", f"-n{limit}"])


async def tool_git_diff(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil git reserve aux administrateurs"
    if not _git_repo_available():
        return _build_metadata() + "\n\nDiff git indisponible car .git n'est pas copie dans l'image Docker."
    try:
        path = _safe_git_pathspec(str(args.get("path", "")))
    except ValueError as e:
        return str(e)
    cmd = ["diff", "--no-ext-diff"]
    if bool(args.get("stat")):
        cmd.append("--stat")
    if path:
        cmd.extend(["--", path])
    return await _run_git(cmd, timeout=15)


async def tool_git_show(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil git reserve aux administrateurs"
    if not _git_repo_available():
        return _build_metadata() + "\n\nDetails commit indisponibles car .git n'est pas copie dans l'image Docker."
    try:
        ref = _safe_git_ref(str(args.get("ref", "HEAD")))
    except ValueError as e:
        return str(e)
    mode = str(args.get("mode", "summary")).strip().lower()
    if mode == "patch":
        cmd = ["show", "--no-ext-diff", "--stat", "--patch", ref]
    else:
        cmd = ["show", "--no-ext-diff", "--stat", "--oneline", "--decorate", ref]
    return await _run_git(cmd, timeout=15)




# --- Date / heure -------------------------------------------------------------

async def tool_get_datetime(args: dict, ctx: dict) -> str:
    now = datetime.datetime.now().astimezone()
    jours = ["lundi","mardi","mercredi","jeudi","vendredi","samedi","dimanche"]
    return (f"{jours[now.weekday()]} {now.strftime('%d/%m/%Y %H:%M:%S')} "
            f"(fuseau {now.tzname() or now.strftime('%z')})")


# --- Requete HTTP generique (anti-SSRF) ----------------------------------------

def _is_private_host(host: str) -> bool:
    try:
        infos = socket.getaddrinfo(host, None)
    except socket.gaierror:
        return True  # irresoluble = refuse
    for info in infos:
        try:
            ip = ipaddress.ip_address(info[4][0])
        except ValueError:
            return True
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
            return True
    return False


MAX_REDIRECTS = 5


def _host_of(url: str) -> str:
    try:
        return (urlparse(url).hostname or "").strip()
    except ValueError:
        return ""


async def _safe_fetch(
    method: str,
    url: str,
    ctx: dict | None = None,
    *,
    timeout: float = 25,
    headers: dict | None = None,
    params: dict | None = None,
    content=None,
    data=None,
    user_agent: str = "Mozilla/5.0 (compatible; Maltai/1.0)",
):
    """Requete HTTP qui suit les redirections MANUELLEMENT en revalidant l'hote
    a chaque saut. Pour les comptes non-admin, tout saut (y compris l'URL
    initiale) vers une IP privee/interne est refuse -> protege contre le SSRF
    par redirection. Leve PermissionError si bloque."""
    is_admin = bool((ctx or {}).get("is_admin"))
    hdrs = {"User-Agent": user_agent}
    if headers:
        hdrs.update(headers)
    current = url
    async with httpx.AsyncClient(timeout=timeout, follow_redirects=False) as client:
        for _ in range(MAX_REDIRECTS + 1):
            if not is_admin and _is_private_host(_host_of(current)):
                raise PermissionError("hote prive/interne refuse (SSRF)")
            resp = await client.request(
                method, current, headers=hdrs,
                params=params, content=content, data=data,
            )
            if resp.is_redirect and resp.headers.get("location"):
                current = urljoin(current, resp.headers["location"])
                # On ne rejoue ni le corps ni les parametres vers la cible suivante.
                params = None
                content = None
                data = None
                method = "GET"
                continue
            return resp
    raise PermissionError("trop de redirections")


async def tool_http_request(args: dict, ctx: dict) -> str:
    url = str(args.get("url", ""))
    method = str(args.get("method", "GET")).upper()
    if method not in ("GET", "POST", "PUT", "PATCH", "DELETE", "HEAD"):
        return "Methode non autorisee"
    if not url.startswith(("http://", "https://")):
        return "URL invalide"
    host = re.sub(r"^https?://", "", url).split("/")[0].split(":")[0]
    if not ctx.get("is_admin") and _is_private_host(host):
        return "Refuse : hote prive/interne (reserve aux administrateurs)"
    headers = args.get("headers") or {}
    if not isinstance(headers, dict):
        headers = {}
    body = args.get("body")
    try:
        r = await _safe_fetch(
            method, url, ctx, timeout=25, headers=headers,
            content=json.dumps(body) if isinstance(body, (dict, list)) else body,
        )
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur HTTP : {e}"
    ctype = r.headers.get("content-type", "")
    text = _strip_html(r.text) if "html" in ctype else r.text
    return _truncate(f"[{r.status_code} {ctype.split(';')[0]}]\n{text}")


# --- Recherche dans la memoire vectorielle --------------------------------------

async def tool_memory_search(args: dict, ctx: dict) -> str:
    from src import memory  # import local pour eviter un cycle
    provider = ctx.get("provider")
    if not provider or not provider.get("embed_model"):
        return "Memoire indisponible (pas de modele d'embeddings sur le provider)"
    query = str(args.get("query", "")).strip()
    if not query:
        return "Requete vide"
    results = await memory.recall(provider, ctx.get("user_id"), query, k=6)
    if not results:
        return "Aucun souvenir pertinent"
    lines = []
    for m in results:
        who = "user" if m["role"] == "user" else "assistant"
        lines.append(f"- ({m['score']:.2f}) [{who}] {m['content'][:300]}")
    return "\n".join(lines)


# --- Execution Python (admin uniquement) ----------------------------------------

async def tool_python_exec(args: dict, ctx: dict) -> str:
    if not ctx.get("is_admin"):
        return "Refuse : outil python reserve aux administrateurs"
    code = str(args.get("code", ""))
    try:
        proc = await asyncio.create_subprocess_exec(
            sys.executable, "-I", "-c", code,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
            cwd=str(WORKSPACE),
        )
        out, _ = await asyncio.wait_for(proc.communicate(), timeout=30)
        return _truncate(f"[exit {proc.returncode}]\n{out.decode(errors='replace')}")
    except asyncio.TimeoutError:
        proc.kill()
        return "Timeout (30s) — code interrompu"
    except OSError as e:
        return f"Erreur python : {e}"



async def tool_code_execute(args: dict, ctx: dict) -> str:
    """Execute du code Python dans un sandbox isole (tous users, timeout 10s).
    Acces reseau et modules systeme bloques, stdout/stderr captures."""
    code = str(args.get("code", "")).strip()
    if not code:
        return "Code vide"
    # Wrapper de securite : bloque les imports dangereux
    guard = (
        "import sys, builtins\n"
        "_blocked = {'os','subprocess','socket','shutil','importlib',"
        "'ctypes','multiprocessing','threading','signal','pty','fcntl','resource'}\n"
        "_orig_import = builtins.__import__\n"
        "def _safe_import(name, *a, **kw):\n"
        "    top = name.split('.')[0]\n"
        "    if top in _blocked:\n"
        "        raise ImportError(f'Import bloque (sandbox): {name}')\n"
        "    return _orig_import(name, *a, **kw)\n"
        "builtins.__import__ = _safe_import\n"
    )
    full_code = guard + "\n" + code
    try:
        import tempfile, os as _os
        with tempfile.TemporaryDirectory() as td:
            code_file = _os.path.join(td, "_code.py")
            with open(code_file, "w") as cf:
                cf.write(full_code)
            proc = await asyncio.create_subprocess_exec(
                sys.executable, "-I", code_file,
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.STDOUT,
                cwd=td,
                env={"PATH": "/usr/bin:/bin"},
            )
            try:
                out, _ = await asyncio.wait_for(proc.communicate(), timeout=10)
            except asyncio.TimeoutError:
                proc.kill()
                return "Timeout (10s) — code interrompu"
        result = out.decode(errors="replace").strip()
        label = f"[exit {proc.returncode}]"
        return _truncate(f"{label}\n{result}" if result else f"{label} (aucune sortie)")
    except OSError as e:
        return f"Erreur sandbox : {e}"


# --- Wikipedia ------------------------------------------------------------------

WIKI_BASE = "https://fr.wikipedia.org"


async def tool_wikipedia(args: dict, ctx: dict) -> str:
    query = str(args.get("query", "")).strip()
    if not query:
        return "Requete vide"
    try:
        async with httpx.AsyncClient(timeout=20) as client:
            s = await client.get(f"{WIKI_BASE}/w/api.php", params={
                "action": "query", "list": "search", "srsearch": query,
                "srlimit": 3, "format": "json",
            })
            s.raise_for_status()
            hits = s.json().get("query", {}).get("search", [])
            if not hits:
                return "Aucun article trouve"
            title = hits[0]["title"]
            r = await client.get(f"{WIKI_BASE}/api/rest_v1/page/summary/{title}")
            r.raise_for_status()
            d = r.json()
    except httpx.HTTPError as e:
        return f"Erreur Wikipedia : {e}"
    out = [f"# {d.get('title', title)}", d.get("extract", ""),
           f"Source : {d.get('content_urls', {}).get('desktop', {}).get('page', '')}"]
    if len(hits) > 1:
        out.append("Autres articles : " + ", ".join(h["title"] for h in hits[1:]))
    return _truncate("\n".join(filter(None, out)))


# --- Meteo (Open-Meteo, sans cle) -----------------------------------------------

GEO_BASE = "https://geocoding-api.open-meteo.com"
METEO_BASE = "https://api.open-meteo.com"

_WMO = {0: "ciel clair", 1: "plutot clair", 2: "partiellement nuageux", 3: "couvert",
        45: "brouillard", 48: "brouillard givrant", 51: "bruine legere", 53: "bruine",
        55: "bruine dense", 61: "pluie legere", 63: "pluie", 65: "pluie forte",
        71: "neige legere", 73: "neige", 75: "neige forte", 80: "averses legeres",
        81: "averses", 82: "averses fortes", 95: "orage", 96: "orage avec grele",
        99: "orage violent avec grele"}


async def tool_weather(args: dict, ctx: dict) -> str:
    city = str(args.get("city", "")).strip()
    if not city:
        return "Ville manquante"
    try:
        async with httpx.AsyncClient(timeout=20) as client:
            g = await client.get(f"{GEO_BASE}/v1/search",
                                 params={"name": city, "count": 1, "language": "fr"})
            g.raise_for_status()
            results = g.json().get("results") or []
            if not results:
                return f"Ville introuvable : {city}"
            loc = results[0]
            f = await client.get(f"{METEO_BASE}/v1/forecast", params={
                "latitude": loc["latitude"], "longitude": loc["longitude"],
                "current": "temperature_2m,weather_code,wind_speed_10m",
                "daily": "temperature_2m_max,temperature_2m_min,precipitation_probability_max,weather_code",
                "timezone": "auto", "forecast_days": 3,
            })
            f.raise_for_status()
            d = f.json()
    except httpx.HTTPError as e:
        return f"Erreur meteo : {e}"
    cur = d.get("current", {})
    daily = d.get("daily", {})
    lines = [
        f"Meteo {loc['name']} ({loc.get('country', '')}) :",
        f"Actuellement : {cur.get('temperature_2m')}°C, "
        f"{_WMO.get(cur.get('weather_code'), 'n/a')}, vent {cur.get('wind_speed_10m')} km/h",
    ]
    for i, day in enumerate(daily.get("time", [])[:3]):
        lines.append(
            f"{day} : {daily['temperature_2m_min'][i]}–{daily['temperature_2m_max'][i]}°C, "
            f"{_WMO.get(daily['weather_code'][i], '')}, "
            f"pluie {daily['precipitation_probability_max'][i]}%"
        )
    return "\n".join(lines)


# --- Flux RSS / Atom ------------------------------------------------------------

async def tool_rss_fetch(args: dict, ctx: dict) -> str:
    url = str(args.get("url", ""))
    if not url.startswith(("http://", "https://")):
        return "URL invalide"
    max_items = min(int(args.get("max_items", 8) or 8), 20)
    try:
        r = await _safe_fetch("GET", url, ctx, timeout=20, user_agent="Maltai/1.1")
        r.raise_for_status()
        root = ET.fromstring(r.content)
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur flux : {e}"
    except ET.ParseError as e:
        return f"Flux illisible : {e}"
    ns = {"atom": "http://www.w3.org/2005/Atom"}
    items = root.findall(".//item")[:max_items] or root.findall(".//atom:entry", ns)[:max_items]
    if not items:
        return "Aucun element dans le flux"
    out = []
    for it in items:
        title = (it.findtext("title") or it.findtext("atom:title", namespaces=ns) or "").strip()
        link = it.findtext("link") or ""
        if not link:
            ln = it.find("atom:link", ns)
            link = ln.get("href") if ln is not None else ""
        date = (it.findtext("pubDate") or it.findtext("atom:updated", namespaces=ns) or "").strip()
        out.append(f"- {title}\n  {link}\n  {date}")
    return _truncate("\n".join(out))


# --- Transcription YouTube -------------------------------------------------------

def _yt_video_id(ref: str) -> str | None:
    ref = ref.strip()
    if re.fullmatch(r"[\w-]{11}", ref):
        return ref
    m = re.search(r"(?:v=|youtu\.be/|shorts/|embed/)([\w-]{11})", ref)
    return m.group(1) if m else None


async def tool_youtube_transcript(args: dict, ctx: dict) -> str:
    vid = _yt_video_id(str(args.get("url_or_id", "")))
    if not vid:
        return "URL/ID YouTube invalide"
    try:
        from youtube_transcript_api import YouTubeTranscriptApi
        loop = asyncio.get_event_loop()
        def fetch():
            api = YouTubeTranscriptApi()
            tr = api.fetch(vid, languages=["fr", "en"])
            return " ".join(seg.text for seg in tr)
        text = await loop.run_in_executor(None, fetch)
    except Exception as e:
        return f"Transcription indisponible : {e}"
    return _truncate(text)


# --- Generation d'images (endpoint compatible OpenAI) ----------------------------

async def tool_generate_image(args: dict, ctx: dict) -> str:
    from core.config import settings
    if not settings.IMAGE_API_BASE:
        return ("Generation d'images non configuree : definis IMAGE_API_BASE "
                "(+ IMAGE_API_KEY / IMAGE_MODEL) vers un endpoint compatible "
                "OpenAI /v1/images/generations (OpenAI, LocalAI, SD-webui, ComfyUI via wrapper...)")
    prompt = str(args.get("prompt", "")).strip()
    if not prompt:
        return "Prompt vide"
    size = str(args.get("size", "1024x1024"))
    base = settings.IMAGE_API_BASE.rstrip("/")
    if not base.endswith("/v1"):
        base += "/v1"
    headers = {"Content-Type": "application/json"}
    if settings.IMAGE_API_KEY:
        headers["Authorization"] = f"Bearer {settings.IMAGE_API_KEY}"
    payload = {"prompt": prompt, "size": size, "response_format": "b64_json", "n": 1}
    if settings.IMAGE_MODEL:
        payload["model"] = settings.IMAGE_MODEL
    try:
        async with httpx.AsyncClient(timeout=180) as client:
            r = await client.post(f"{base}/images/generations", headers=headers, json=payload)
            r.raise_for_status()
            data = r.json().get("data", [])
    except httpx.HTTPError as e:
        return f"Erreur generation image : {e}"
    if not data:
        return "Aucune image retournee"
    b64 = data[0].get("b64_json", "")
    if not b64:
        url = data[0].get("url", "")
        return f"Image generee (URL distante) : {url}" if url else "Reponse sans image"
    img_dir = WORKSPACE / "images"
    img_dir.mkdir(exist_ok=True)
    import uuid as _uuid
    path = img_dir / f"img_{_uuid.uuid4().hex[:10]}.png"
    path.write_bytes(base64.b64decode(b64))
    rel = path.relative_to(WORKSPACE)
    return (f"Image generee : {rel} ({path.stat().st_size // 1024} Ko). "
            f"Telechargeable dans Reglages > Workspace de l'agent.")


# --- Deep research ----------------------------------------------------------------

RESEARCH_MAX_CHARS = 12000


async def tool_deep_research(args: dict, ctx: dict) -> str:
    """Recherche approfondie : plan de requetes -> recherches -> lecture des
    meilleures pages -> rapport markdown synthetise par le modele."""
    from src import llm as _llm
    provider = ctx.get("provider")
    model = ctx.get("model")
    if not provider or not model:
        return "Deep research indisponible (provider/modele manquant)"
    topic = str(args.get("topic", "")).strip()
    if not topic:
        return "Sujet vide"

    async def ask(prompt: str) -> str:
        try:
            return await _llm.complete(
                provider["base_url"], provider["api_key"], model,
                [{"role": "user", "content": prompt}],
                max_tokens=ctx.get("max_tokens"),
            )
        except _llm.LLMError as e:
            raise RuntimeError(str(e))

    try:
        # 1. Plan : 3 requetes de recherche
        plan = await ask(
            f"Sujet de recherche : {topic}\n"
            "Donne exactement 3 requetes de recherche web courtes (3-6 mots), "
            "une par ligne, sans numerotation ni autre texte."
        )
        queries = [q.strip("-• ").strip() for q in plan.splitlines() if q.strip()][:3] or [topic]

        # 2. Recherches
        findings = []
        urls: list[str] = []
        for q in queries:
            res = await tool_web_search({"query": q}, ctx)
            findings.append(f"### Recherche : {q}\n{res}")
            urls += re.findall(r"https?://\S+", res)

        # 3. Lecture des 3 premieres pages distinctes
        seen, picked = set(), []
        for u in urls:
            host = u.split("/")[2] if "//" in u else u
            if host not in seen:
                seen.add(host)
                picked.append(u)
            if len(picked) >= 3:
                break
        pages = []
        for u in picked:
            content = await tool_web_fetch({"url": u}, ctx)
            pages.append(f"### Source : {u}\n{content[:3000]}")

        # 4. Synthese
        corpus = "\n\n".join(findings + pages)[:16000]
        report = await ask(
            f"Sujet : {topic}\n\nDonnees collectees :\n{corpus}\n\n"
            "Redige un rapport structure en markdown (titres ##, points cles, "
            "chiffres si disponibles) en francais, puis une section 'Sources' "
            "listant les URLs utilisees. Sois factuel et concis."
        )
    except RuntimeError as e:
        return f"Erreur deep research : {e}"
    if len(report) > RESEARCH_MAX_CHARS:
        report = report[:RESEARCH_MAX_CHARS] + "\n…[tronque]"
    return report


# --- Notes & taches (inspire d'Odysseus) ------------------------------------

def _fmt_ts(ts: float) -> str:
    return datetime.datetime.fromtimestamp(ts).strftime("%Y-%m-%d %H:%M")


def _find_note(conn, user_id, id_prefix: str, kind: str):
    rows = conn.execute(
        "SELECT * FROM notes WHERE kind=? AND (user_id IS ? OR user_id=?) "
        "AND id LIKE ? ORDER BY created_at",
        (kind, user_id, user_id, id_prefix + "%"),
    ).fetchall()
    return rows


async def tool_note_add(args: dict, ctx: dict) -> str:
    from core import database as db
    content = str(args.get("content", "")).strip()
    if not content:
        return "Contenu vide"
    conn = db.connect()
    try:
        nid = db.new_id()
        conn.execute(
            "INSERT INTO notes (id, user_id, kind, content, done, created_at) "
            "VALUES (?, ?, 'note', ?, 0, ?)",
            (nid, ctx.get("user_id"), content, db.now()),
        )
        conn.commit()
        return f"Note enregistree (id {nid[:8]})"
    finally:
        conn.close()


async def tool_note_list(args: dict, ctx: dict) -> str:
    from core import database as db
    conn = db.connect()
    try:
        rows = conn.execute(
            "SELECT * FROM notes WHERE kind='note' AND (user_id IS ? OR user_id=?) "
            "ORDER BY created_at DESC LIMIT 50",
            (ctx.get("user_id"), ctx.get("user_id")),
        ).fetchall()
    finally:
        conn.close()
    if not rows:
        return "Aucune note"
    lines = [f"- [{r['id'][:8]}] ({_fmt_ts(r['created_at'])}) {r['content']}" for r in rows]
    return _truncate("\n".join(lines))


async def tool_note_delete(args: dict, ctx: dict) -> str:
    from core import database as db
    id_prefix = str(args.get("note_id", "")).strip()
    if not id_prefix:
        return "note_id manquant"
    conn = db.connect()
    try:
        rows = _find_note(conn, ctx.get("user_id"), id_prefix, "note")
        if not rows:
            return f"Aucune note avec l'id {id_prefix}"
        if len(rows) > 1:
            return f"Id ambigu ({len(rows)} notes). Precise plus de caracteres."
        conn.execute("DELETE FROM notes WHERE id=?", (rows[0]["id"],))
        conn.commit()
        return f"Note supprimee : {rows[0]['content'][:80]}"
    finally:
        conn.close()


async def tool_todo_add(args: dict, ctx: dict) -> str:
    from core import database as db
    content = str(args.get("content", "")).strip()
    if not content:
        return "Contenu vide"
    conn = db.connect()
    try:
        nid = db.new_id()
        conn.execute(
            "INSERT INTO notes (id, user_id, kind, content, done, created_at) "
            "VALUES (?, ?, 'todo', ?, 0, ?)",
            (nid, ctx.get("user_id"), content, db.now()),
        )
        conn.commit()
        return f"Tache ajoutee (id {nid[:8]})"
    finally:
        conn.close()


async def tool_todo_list(args: dict, ctx: dict) -> str:
    from core import database as db
    include_done = bool(args.get("include_done", False))
    conn = db.connect()
    try:
        sql = ("SELECT * FROM notes WHERE kind='todo' AND (user_id IS ? OR user_id=?) "
               + ("" if include_done else "AND done=0 ")
               + "ORDER BY done, created_at DESC LIMIT 100")
        rows = conn.execute(sql, (ctx.get("user_id"), ctx.get("user_id"))).fetchall()
    finally:
        conn.close()
    if not rows:
        return "Aucune tache" + (" (toutes terminees ?)" if not include_done else "")
    lines = []
    for r in rows:
        box = "[x]" if r["done"] else "[ ]"
        lines.append(f"- {box} [{r['id'][:8]}] {r['content']}")
    return _truncate("\n".join(lines))


async def tool_todo_done(args: dict, ctx: dict) -> str:
    from core import database as db
    id_prefix = str(args.get("todo_id", "")).strip()
    if not id_prefix:
        return "todo_id manquant"
    conn = db.connect()
    try:
        rows = _find_note(conn, ctx.get("user_id"), id_prefix, "todo")
        if not rows:
            return f"Aucune tache avec l'id {id_prefix}"
        if len(rows) > 1:
            return f"Id ambigu ({len(rows)} taches). Precise plus de caracteres."
        conn.execute("UPDATE notes SET done=1 WHERE id=?", (rows[0]["id"],))
        conn.commit()
        return f"Tache terminee : {rows[0]['content'][:80]}"
    finally:
        conn.close()


# --- Envoi d'email (SMTP, inspire d'Odysseus) --------------------------------

_EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")


async def tool_email_send(args: dict, ctx: dict) -> str:
    from core.config import settings
    if not settings.SMTP_HOST or not settings.SMTP_USER:
        return ("Email non configure. Definis les variables d'environnement "
                "SMTP_HOST, SMTP_PORT, SMTP_USER, SMTP_PASSWORD (et SMTP_FROM optionnel).")
    to = str(args.get("to", "")).strip()
    subject = str(args.get("subject", "")).strip()
    body = str(args.get("body", ""))
    if not _EMAIL_RE.match(to):
        return f"Adresse destinataire invalide : {to}"
    if not subject or not body:
        return "subject et body sont requis"

    import smtplib
    from email.message import EmailMessage

    msg = EmailMessage()
    msg["From"] = settings.SMTP_FROM or settings.SMTP_USER
    msg["To"] = to
    msg["Subject"] = subject
    msg.set_content(body)

    def send():
        with smtplib.SMTP(settings.SMTP_HOST, settings.SMTP_PORT, timeout=20) as s:
            if settings.SMTP_TLS:
                s.starttls()
            s.login(settings.SMTP_USER, settings.SMTP_PASSWORD)
            s.send_message(msg)

    try:
        await asyncio.to_thread(send)
        return f"Email envoye a {to} (sujet : {subject})"
    except Exception as e:
        return f"Echec de l'envoi : {e}"


# =============================================================================
# Pack v1.3 — memory_save, session_search, patch_file, skill_save/list/run
# =============================================================================

async def tool_memory_save(args: dict, ctx: dict) -> str:
    """Ecrit un fait dans la memoire vectorielle persistante de l'utilisateur."""
    from core import database as db
    from src.memory import embed_text, cosine_sim  # noqa: F401
    import struct, numpy as np

    fact = str(args.get("fact", "")).strip()
    if not fact:
        return "Erreur : parametre 'fact' manquant."
    user_id = ctx.get("user_id")
    session_id = ctx.get("session_id")
    provider = ctx.get("provider")
    if provider is None:
        return "Pas de provider configure, impossible de calculer l'embedding."
    try:
        emb = await embed_text(fact, provider)
    except Exception as e:
        return f"Erreur embedding : {e}"
    dim = len(emb)
    blob = struct.pack(f"{dim}f", *emb)
    mid = db.add_memory(user_id, session_id, "user", fact, blob, dim)
    return f"Souvenir enregistre (id={mid[:8]}…) : « {fact[:80]} »"


async def tool_session_search(args: dict, ctx: dict) -> str:
    """Recherche plein texte dans toutes les conversations (FTS5)."""
    from core import database as db

    query = str(args.get("query", "")).strip()
    if not query:
        return "Erreur : parametre 'query' manquant."
    limit = min(int(args.get("limit", 8)), 20)
    user_id = ctx.get("user_id")
    results = db.fts_search(query, user_id, limit)
    if not results:
        return f"Aucun resultat pour « {query} »."
    if results and "error" in results[0]:
        return f"Erreur FTS : {results[0]['error']}"
    lines = [f"Resultats pour « {query} » ({len(results)}) :\n"]
    for r in results:
        import datetime
        ts = datetime.datetime.fromtimestamp(r["updated_at"]).strftime("%d/%m/%Y")
        lines.append(f"- [{r['session_title']}] ({ts}) [{r['role']}] {r['snippet']}")
    return "\n".join(lines)


async def tool_patch_file(args: dict, ctx: dict) -> str:
    """Remplace un bloc de texte dans un fichier workspace par un nouveau contenu."""
    path_str = str(args.get("path", "")).strip()
    old_text = args.get("old_str", "")
    new_text = args.get("new_str", "")
    if not path_str:
        return "Erreur : parametre 'path' manquant."
    try:
        target = _safe_path(path_str, ctx)
    except ValueError as e:
        return str(e)
    if not target.exists():
        return f"Fichier introuvable : {path_str}"
    content = target.read_text(encoding="utf-8")
    if old_text and old_text not in content:
        return "Erreur : le bloc 'old_str' est introuvable dans le fichier."
    if old_text:
        content = content.replace(old_text, new_text, 1)
    else:
        content = new_text
    target.write_text(content, encoding="utf-8")
    return f"Fichier mis a jour : {path_str} ({len(new_text)} caracteres inseres)"


async def tool_skill_save(args: dict, ctx: dict) -> str:
    """Sauvegarde une procedure reutilisable (skill) dans la base."""
    from core import database as db

    name = str(args.get("name", "")).strip().lower().replace(" ", "_")
    description = str(args.get("description", "")).strip()
    body = str(args.get("body", "")).strip()
    if not name or not body:
        return "Erreur : 'name' et 'body' sont requis."
    user_id = ctx.get("user_id")
    db.skill_save(user_id, name, description, body)
    return f"Skill « {name} » sauvegarde."


async def tool_skill_list(args: dict, ctx: dict) -> str:
    """Liste les skills disponibles pour cet utilisateur."""
    from core import database as db

    user_id = ctx.get("user_id")
    skills = db.list_skills(user_id)
    if not skills:
        return "Aucun skill enregistre. Utilisez skill_save pour en creer un."
    lines = ["Skills disponibles :\n"]
    for s in skills:
        desc = f" — {s['description']}" if s['description'] else ""
        lines.append(f"• **{s['name']}**{desc}")
    return "\n".join(lines)


async def tool_skill_run(args: dict, ctx: dict) -> str:
    """Rappelle le corps d'un skill et le retourne pour que l'agent l'execute."""
    from core import database as db

    name = str(args.get("name", "")).strip().lower().replace(" ", "_")
    if not name:
        return "Erreur : 'name' est requis."
    user_id = ctx.get("user_id")
    skill = db.get_skill(user_id, name)
    if not skill:
        return f"Skill « {name} » introuvable. Verifiez le nom avec skill_list."
    return f"=== Skill : {skill['name']} ===\n{skill['body']}"


# =============================================================================
# Pack v1.4 — image_generate, page_summary
# =============================================================================

async def tool_image_generate(args: dict, ctx: dict) -> str:
    """Genere une image via l'API images du provider OpenAI-compatible."""
    provider = ctx.get("provider")
    if not provider:
        return "Erreur : aucun provider configure."
    prompt = str(args.get("prompt", "")).strip()
    if not prompt:
        return "Erreur : parametre 'prompt' manquant."
    size = str(args.get("size", "1024x1024"))
    quality = str(args.get("quality", "standard"))
    n = int(args.get("n", 1))
    base_url = provider["base_url"].rstrip("/")
    # Remove /v1 suffix if already present to reconstruct cleanly
    if base_url.endswith("/v1"):
        base_url = base_url[:-3]
    url = f"{base_url}/v1/images/generations"
    headers = {"Content-Type": "application/json"}
    if provider.get("api_key"):
        headers["Authorization"] = f"Bearer {provider['api_key']}"
    payload = {"prompt": prompt, "n": n, "size": size}
    # quality param only supported by some providers (dall-e-3)
    if quality != "standard":
        payload["quality"] = quality
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            r = await client.post(url, json=payload, headers=headers)
        if r.status_code != 200:
            return f"Erreur API images ({r.status_code}) : {r.text[:300]}"
        data = r.json()
        images = data.get("data", [])
        if not images:
            return "Aucune image retournee par le provider."
        urls = []
        for img in images:
            if img.get("url"):
                urls.append(img["url"])
            elif img.get("b64_json"):
                urls.append(f"data:image/png;base64,{img['b64_json'][:40]}…")
        return "Image(s) generee(s) :\n" + "\n".join(urls)
    except httpx.HTTPError as e:
        return f"Erreur reseau : {e}"


async def tool_page_summary(args: dict, ctx: dict) -> str:
    """Recupere le contenu d'une page web et en extrait le texte principal."""
    import re as _re
    url = str(args.get("url", "")).strip()
    if not url or not url.startswith(("http://", "https://")):
        return "Erreur : URL invalide ou manquante."
    question = str(args.get("question", "")).strip()
    try:
        r = await _safe_fetch(
            "GET", url, ctx, timeout=20,
            user_agent="Mozilla/5.0 (compatible; MaltaiBot/1.0)",
        )
        if r.status_code != 200:
            return f"Erreur HTTP {r.status_code} sur {url}"
        raw = r.text
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur reseau : {e}"

    # Strip scripts, styles, HTML tags
    raw = _re.sub(r"<(script|style|head)[^>]*>.*?</(script|style|head)>", " ", raw, flags=_re.S | _re.I)
    raw = _re.sub(r"<[^>]+>", " ", raw)
    raw = html.unescape(raw)
    raw = _re.sub(r"[ \t]{2,}", " ", raw)
    raw = _re.sub(r"\n{3,}", "\n\n", raw).strip()

    # Truncate to ~6000 chars for context
    excerpt = raw[:6000]
    if len(raw) > 6000:
        excerpt += f"\n\n[... {len(raw)-6000} caracteres supprimes ...]"

    if question:
        return f"Page : {url}\nQuestion : {question}\n\nContenu :\n{excerpt}"
    return f"Page : {url}\n\nContenu extrait :\n{excerpt}"


def _same_site(url: str, base_host: str) -> bool:
    try:
        host = urlparse(url).hostname or ""
    except ValueError:
        return False
    return host.lower() == base_host.lower()


async def tool_web_crawl(args: dict, ctx: dict) -> str:
    """Explore quelques pages d'un meme site et extrait titres/liens/resumes."""
    start_url = str(args.get("url", "")).strip()
    if not start_url.startswith(("http://", "https://")):
        return "URL invalide"
    try:
        base_host = urlparse(start_url).hostname or ""
    except ValueError:
        return "URL invalide"
    max_pages = max(1, min(40, int(args.get("max_pages") or 10)))
    max_depth = max(0, min(3, int(args.get("max_depth") or 1)))
    include_text = bool(args.get("include_text", False))
    save_as = str(args.get("save_as", "") or "").strip()

    seen: set[str] = set()
    queue: list[tuple[str, int]] = [(start_url, 0)]
    pages = []
    try:
        while queue and len(pages) < max_pages:
            current, depth = queue.pop(0)
            current = current.split("#")[0]
            if current in seen:
                continue
            seen.add(current)
            r = await _safe_fetch("GET", current, ctx, timeout=20, user_agent="Mozilla/5.0 (compatible; MaltaiCrawler/1.0)")
            ctype = r.headers.get("content-type", "")
            if r.status_code >= 400 or "html" not in ctype:
                pages.append({"url": str(r.url), "status": r.status_code, "content_type": ctype})
                continue
            raw = r.text
            text = _strip_html(raw)
            links = []
            for link in _html_links(raw, str(r.url), limit=120):
                href = link.get("url", "").split("#")[0]
                if href.startswith(("http://", "https://")) and _same_site(href, base_host):
                    links.append(href)
            pages.append({
                "url": str(r.url),
                "status": r.status_code,
                "title": _html_title(raw),
                "h1": re.findall(r"<h1[^>]*>(.*?)</h1>", raw, re.I | re.S)[:3],
                "links_found": len(set(links)),
                **({"text": text[:2500]} if include_text else {"excerpt": text[:500]}),
            })
            if depth < max_depth:
                for href in links:
                    if href not in seen and all(href != q[0] for q in queue):
                        queue.append((href, depth + 1))
        result = {"start_url": start_url, "pages_crawled": len(pages), "pages": pages}
        if save_as:
            target = _safe_path(save_as if save_as.lower().endswith(".json") else save_as + ".json", ctx)
            target.parent.mkdir(parents=True, exist_ok=True)
            target.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
            rel = str(target.relative_to(_user_workspace(ctx))).replace("\\", "/")
            result["export"] = {"path": rel, "download_url": f"/api/workspace/download?path={quote(rel)}"}
        return _truncate(json.dumps(result, ensure_ascii=False, indent=2))
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur crawl : {e}"


async def tool_seo_audit(args: dict, ctx: dict) -> str:
    """Audit SEO technique rapide d'une page HTML."""
    url = str(args.get("url", "")).strip()
    if not url.startswith(("http://", "https://")):
        return "URL invalide"
    try:
        r = await _safe_fetch("GET", url, ctx, timeout=20, user_agent="Mozilla/5.0 (compatible; MaltaiSEO/1.0)")
        if r.status_code >= 400:
            return f"Erreur HTTP {r.status_code} sur {url}"
        raw = r.text
    except PermissionError as e:
        return f"Refuse : {e}"
    except httpx.HTTPError as e:
        return f"Erreur SEO : {e}"

    title = _html_title(raw)
    meta_desc = ""
    m = re.search(r'<meta[^>]+name=["\']description["\'][^>]*content=["\']([^"\']*)', raw, re.I)
    if not m:
        m = re.search(r'<meta[^>]+content=["\']([^"\']*)["\'][^>]+name=["\']description["\']', raw, re.I)
    if m:
        meta_desc = html.unescape(m.group(1).strip())
    h1s = [_strip_html(x) for x in re.findall(r"<h1[^>]*>(.*?)</h1>", raw, re.I | re.S)]
    h2s = [_strip_html(x) for x in re.findall(r"<h2[^>]*>(.*?)</h2>", raw, re.I | re.S)]
    imgs = re.findall(r"<img\b[^>]*>", raw, re.I)
    imgs_missing_alt = [tag[:180] for tag in imgs if not re.search(r"\salt\s*=", tag, re.I)]
    links = _html_links(raw, str(r.url), limit=300)
    canonical = ""
    cm = re.search(r'<link[^>]+rel=["\']canonical["\'][^>]*href=["\']([^"\']+)', raw, re.I)
    if cm:
        canonical = urljoin(str(r.url), cm.group(1))
    viewport = bool(re.search(r'<meta[^>]+name=["\']viewport["\']', raw, re.I))
    robots_noindex = bool(re.search(r'<meta[^>]+name=["\']robots["\'][^>]+content=["\'][^"\']*noindex', raw, re.I))
    issues = []
    if not title:
        issues.append("title manquant")
    elif len(title) < 20 or len(title) > 65:
        issues.append(f"title longueur a verifier ({len(title)} caracteres)")
    if not meta_desc:
        issues.append("meta description manquante")
    elif len(meta_desc) < 70 or len(meta_desc) > 170:
        issues.append(f"meta description longueur a verifier ({len(meta_desc)} caracteres)")
    if len(h1s) != 1:
        issues.append(f"H1 attendu: 1, trouve: {len(h1s)}")
    if imgs_missing_alt:
        issues.append(f"{len(imgs_missing_alt)} image(s) sans alt")
    if not canonical:
        issues.append("canonical manquant")
    if not viewport:
        issues.append("viewport mobile manquant")
    if robots_noindex:
        issues.append("page en noindex")

    result = {
        "url": str(r.url),
        "status": r.status_code,
        "title": title,
        "title_length": len(title),
        "description": meta_desc,
        "description_length": len(meta_desc),
        "canonical": canonical,
        "h1": h1s,
        "h2_count": len(h2s),
        "images": len(imgs),
        "images_missing_alt": len(imgs_missing_alt),
        "links_count": len(links),
        "viewport": viewport,
        "robots_noindex": robots_noindex,
        "issues": issues,
        "score": max(0, 100 - len(issues) * 12),
    }
    save_as = str(args.get("save_as", "") or "").strip()
    if save_as:
        target = _safe_path(save_as if save_as.lower().endswith(".json") else save_as + ".json", ctx)
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")
        rel = str(target.relative_to(_user_workspace(ctx))).replace("\\", "/")
        result["export"] = {"path": rel, "download_url": f"/api/workspace/download?path={quote(rel)}"}
    return _truncate(json.dumps(result, ensure_ascii=False, indent=2))

# --- Registre ----------------------------------------------------------------

Tool = dict[str, Any]

TOOLS: dict[str, dict] = {
    "calculator": {
        "run": tool_calculator,
        "spec": {
            "name": "calculator",
            "description": "Evalue une expression arithmetique (+ - * / // % **).",
            "parameters": {
                "type": "object",
                "properties": {"expression": {"type": "string", "description": "Ex: (12*7)+3**2"}},
                "required": ["expression"],
            },
        },
    },
    "list_files": {
        "run": tool_list_files,
        "spec": {
            "name": "list_files",
            "description": "Liste les fichiers du workspace de l'agent.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string", "description": "Sous-dossier (defaut: racine)"}},
            },
        },
    },
    "read_file": {
        "run": tool_read_file,
        "spec": {
            "name": "read_file",
            "description": "Lit un fichier texte du workspace.",
            "parameters": {
                "type": "object",
                "properties": {"path": {"type": "string"}},
                "required": ["path"],
            },
        },
    },
    "write_file": {
        "run": tool_write_file,
        "spec": {
            "name": "write_file",
            "description": "Ecrit (ou ecrase) un fichier texte dans le workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "content": {"type": "string"},
                },
                "required": ["path", "content"],
            },
        },
    },
    "pdf_read": {
        "run": tool_pdf_read,
        "spec": {
            "name": "pdf_read",
            "description": "Lit un PDF du workspace et extrait son texte, ses metadonnees et le nombre de pages.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin du PDF dans le workspace"},
                    "start_page": {"type": "integer", "description": "Page de debut, 1 par defaut"},
                    "end_page": {"type": "integer", "description": "Page de fin optionnelle"},
                    "max_chars": {"type": "integer", "description": "Limite de texte retourne, defaut 12000"},
                    "password": {"type": "string", "description": "Mot de passe si le PDF est chiffre"},
                },
                "required": ["path"],
            },
        },
    },
    "pdf_create": {
        "run": tool_pdf_create,
        "spec": {
            "name": "pdf_create",
            "description": "Cree un PDF telechargeable dans le workspace depuis du texte ou un fichier texte/Markdown simple.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin du PDF a creer, ex: exports/rapport.pdf"},
                    "title": {"type": "string", "description": "Titre affiche en haut du PDF"},
                    "content": {"type": "string", "description": "Texte ou Markdown simple (#, ##, listes)"},
                    "source_path": {"type": "string", "description": "Fichier texte source du workspace si content est vide"},
                    "author": {"type": "string", "description": "Auteur du PDF"},
                    "page_size": {"type": "string", "description": "A4 ou letter"},
                },
            },
        },
    },
    "docx_read": {
        "run": tool_docx_read,
        "spec": {
            "name": "docx_read",
            "description": "Lit un document Word DOCX du workspace : paragraphes, tableaux et metadonnees.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin du .docx dans le workspace"},
                    "max_chars": {"type": "integer", "description": "Limite de texte retourne, defaut 12000"},
                },
                "required": ["path"],
            },
        },
    },
    "docx_create": {
        "run": tool_docx_create,
        "spec": {
            "name": "docx_create",
            "description": "Cree un document Word DOCX depuis du texte ou Markdown simple (#, ##, listes).",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin du .docx a creer, ex: exports/rapport.docx"},
                    "title": {"type": "string"},
                    "content": {"type": "string"},
                    "source_path": {"type": "string", "description": "Fichier texte source si content est vide"},
                    "author": {"type": "string"},
                },
            },
        },
    },
    "xlsx_read": {
        "run": tool_xlsx_read,
        "spec": {
            "name": "xlsx_read",
            "description": "Lit un fichier Excel XLSX du workspace et retourne les lignes d'une feuille.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string"},
                    "sheet": {"type": "string", "description": "Nom de feuille optionnel"},
                    "max_rows": {"type": "integer", "description": "Defaut 80, max 500"},
                    "max_cols": {"type": "integer", "description": "Defaut 20, max 80"},
                },
                "required": ["path"],
            },
        },
    },
    "xlsx_create": {
        "run": tool_xlsx_create,
        "spec": {
            "name": "xlsx_create",
            "description": "Cree un fichier Excel XLSX depuis rows, csv ou une liste d'objets JSON.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin du .xlsx a creer, ex: exports/data.xlsx"},
                    "sheet": {"type": "string"},
                    "rows": {"type": "array", "description": "Liste de lignes, ex: [[\"Nom\",\"Prix\"],[\"A\",10]]"},
                    "csv": {"type": "string", "description": "CSV brut optionnel"},
                    "headers": {"type": "array", "description": "Colonnes pour data"},
                    "data": {"type": "array", "description": "Liste d'objets JSON ou lignes"},
                    "header": {"type": "boolean", "description": "Styliser la premiere ligne, defaut true"},
                },
            },
        },
    },
    "zip_create": {
        "run": tool_zip_create,
        "spec": {
            "name": "zip_create",
            "description": "Cree une archive ZIP d'un dossier ou d'une liste de fichiers du workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Archive a creer, ex: exports/workspace.zip"},
                    "folder": {"type": "string", "description": "Dossier a zipper si files est vide"},
                    "files": {"type": "array", "description": "Liste de fichiers/dossiers a inclure"},
                    "include_hidden": {"type": "boolean"},
                    "max_files": {"type": "integer"},
                },
            },
        },
    },
    "zip_extract": {
        "run": tool_zip_extract,
        "spec": {
            "name": "zip_extract",
            "description": "Extrait une archive ZIP dans un dossier du workspace, avec protection anti path traversal.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Archive ZIP du workspace"},
                    "dest": {"type": "string", "description": "Dossier destination, defaut extracted"},
                    "overwrite": {"type": "boolean"},
                    "max_files": {"type": "integer"},
                },
                "required": ["path"],
            },
        },
    },
    "context_compress": {
        "run": tool_context_compress,
        "spec": {
            "name": "context_compress",
            "description": "Compresse un gros texte, JSON, log ou fichier workspace pour reduire les tokens avant analyse.",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "Texte brut optionnel a compresser"},
                    "path": {"type": "string", "description": "Fichier du workspace a compresser si text est vide"},
                    "mode": {"type": "string", "description": "auto, json, log, text ou code"},
                    "max_chars": {"type": "integer", "description": "Taille cible approximative, defaut 3500"},
                },
            },
        },
    },
    "web_search": {
        "run": tool_web_search,
        "spec": {
            "name": "web_search",
            "description": "Recherche web (DuckDuckGo). Retourne titres, URLs et extraits.",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string"}},
                "required": ["query"],
            },
        },
    },
    "browser_navigate": {
        "run": tool_browser_navigate,
        "spec": {
            "name": "browser_navigate",
            "description": "Ouvre une page web dans le navigateur texte de Maltai et garde son etat pour les outils browser_*.",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string", "description": "URL http/https a ouvrir"}},
                "required": ["url"],
            },
        },
    },
    "browser_snapshot": {
        "run": tool_browser_snapshot,
        "spec": {
            "name": "browser_snapshot",
            "description": "Retourne un snapshot lisible de la page ouverte : titre, texte, liens et formulaires.",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string", "description": "URL optionnelle a ouvrir avant le snapshot"}},
            },
        },
    },
    "browser_links": {
        "run": tool_browser_links,
        "spec": {
            "name": "browser_links",
            "description": "Liste les liens trouves sur la page actuellement ouverte par browser_navigate.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    "browser_form_list": {
        "run": tool_browser_form_list,
        "spec": {
            "name": "browser_form_list",
            "description": "Liste les formulaires de la page ouverte : index, methode, action et champs.",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string", "description": "URL optionnelle a ouvrir avant de lister les formulaires"}},
            },
        },
    },
    "browser_submit": {
        "run": tool_browser_submit,
        "spec": {
            "name": "browser_submit",
            "description": "Soumet un formulaire simple de la page ouverte via GET ou POST. N'execute pas de JavaScript.",
            "parameters": {
                "type": "object",
                "properties": {
                    "index": {"type": "integer", "description": "Index du formulaire, defaut 0"},
                    "data": {"type": "object", "description": "Champs a envoyer"},
                    "method": {"type": "string", "description": "GET ou POST optionnel"},
                    "action": {"type": "string", "description": "URL action optionnelle"},
                },
            },
        },
    },
    "browser_open": {
        "run": tool_browser_open,
        "spec": {
            "name": "browser_open",
            "description": "Ouvre une page dans un vrai navigateur Chromium headless via Playwright.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL http/https a ouvrir"},
                    "wait_until": {"type": "string", "description": "load | domcontentloaded | networkidle"},
                    "timeout_ms": {"type": "integer"},
                },
                "required": ["url"],
            },
        },
    },
    "browser_click": {
        "run": tool_browser_click,
        "spec": {
            "name": "browser_click",
            "description": "Clique dans la page ouverte avec Playwright, par selecteur CSS ou texte visible.",
            "parameters": {
                "type": "object",
                "properties": {
                    "selector": {"type": "string", "description": "Selecteur CSS optionnel"},
                    "text": {"type": "string", "description": "Texte visible optionnel"},
                    "timeout_ms": {"type": "integer"},
                },
            },
        },
    },
    "browser_type": {
        "run": tool_browser_type,
        "spec": {
            "name": "browser_type",
            "description": "Saisit du texte dans un champ via Playwright, par selecteur CSS ou label.",
            "parameters": {
                "type": "object",
                "properties": {
                    "selector": {"type": "string", "description": "Selecteur CSS optionnel"},
                    "label": {"type": "string", "description": "Label accessible optionnel"},
                    "text": {"type": "string"},
                    "clear": {"type": "boolean"},
                    "timeout_ms": {"type": "integer"},
                },
                "required": ["text"],
            },
        },
    },
    "browser_screenshot": {
        "run": tool_browser_screenshot,
        "spec": {
            "name": "browser_screenshot",
            "description": "Prend une capture PNG de la page Playwright ouverte et la sauvegarde dans le workspace.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Nom du fichier PNG optionnel"},
                    "full_page": {"type": "boolean"},
                },
            },
        },
    },
    "web_fetch": {
        "run": tool_web_fetch,
        "spec": {
            "name": "web_fetch",
            "description": "Recupere le contenu texte d'une page web (URL exacte).",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string"}},
                "required": ["url"],
            },
        },
    },
    "web_scrape": {
        "run": tool_web_scrape,
        "spec": {
            "name": "web_scrape",
            "description": "Scrape une page web et extrait des donnees structurees JSON : metadata, titres, liens, images, tables, JSON-LD et champs via selecteurs CSS simples.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL http/https a scraper"},
                    "fields": {
                        "type": "object",
                        "description": "Champs a extraire. Ex: {\"titre\":\"h1\", \"prix\":{\"selector\":\".price\", \"attr\":\"text\"}, \"liens\":{\"selector\":\"a\", \"attr\":\"href\", \"all\":true}}",
                    },
                    "include": {
                        "type": "object",
                        "description": "Options booleennes: metadata, headings, links, images, tables, json_ld, text",
                    },
                    "save_as": {"type": "string", "description": "Nom du fichier a creer dans exports/ (ex: arf-reparation.json, data.csv, rapport.md, page.txt, rapport.html)"},
                    "format": {"type": "string", "description": "Format d'export: json, csv, md, txt ou html"},
                    "limit": {"type": "integer", "description": "Nombre maximum d'elements par liste, max 100"},
                },
                "required": ["url"],
            },
        },
    },
    "web_crawl": {
        "run": tool_web_crawl,
        "spec": {
            "name": "web_crawl",
            "description": "Explore plusieurs pages d'un meme site et retourne titres, extraits et liens trouves.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL de depart"},
                    "max_pages": {"type": "integer", "description": "Defaut 10, max 40"},
                    "max_depth": {"type": "integer", "description": "Defaut 1, max 3"},
                    "include_text": {"type": "boolean"},
                    "save_as": {"type": "string", "description": "Exporter le resultat JSON dans le workspace"},
                },
                "required": ["url"],
            },
        },
    },
    "seo_audit": {
        "run": tool_seo_audit,
        "spec": {
            "name": "seo_audit",
            "description": "Audit SEO rapide d'une page : title, description, H1, canonical, images alt, viewport, liens et score.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string"},
                    "save_as": {"type": "string", "description": "Exporter l'audit JSON dans le workspace"},
                },
                "required": ["url"],
            },
        },
    },
    "get_datetime": {
        "run": tool_get_datetime,
        "spec": {
            "name": "get_datetime",
            "description": "Donne la date et l'heure actuelles (fuseau du serveur).",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    "git_status": {
        "run": tool_git_status,
        "spec": {
            "name": "git_status",
            "description": "Affiche l'etat git read-only de l'installation Maltai (ADMIN seulement).",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    "git_branch": {
        "run": tool_git_branch,
        "spec": {
            "name": "git_branch",
            "description": "Affiche branche, commit courant et remote git (ADMIN seulement).",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    "git_log": {
        "run": tool_git_log,
        "spec": {
            "name": "git_log",
            "description": "Liste les derniers commits git read-only (ADMIN seulement).",
            "parameters": {
                "type": "object",
                "properties": {"limit": {"type": "integer", "description": "Nombre de commits, max 30"}},
            },
        },
    },
    "git_diff": {
        "run": tool_git_diff,
        "spec": {
            "name": "git_diff",
            "description": "Affiche le diff git read-only, optionnellement limite a un fichier (ADMIN seulement).",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin relatif optionnel"},
                    "stat": {"type": "boolean", "description": "Afficher seulement les statistiques"},
                },
            },
        },
    },
    "git_show": {
        "run": tool_git_show,
        "spec": {
            "name": "git_show",
            "description": "Affiche un commit git en lecture seule (ADMIN seulement).",
            "parameters": {
                "type": "object",
                "properties": {
                    "ref": {"type": "string", "description": "Reference git, defaut HEAD"},
                    "mode": {"type": "string", "description": "summary ou patch"},
                },
            },
        },
    },
    "http_request": {
        "run": tool_http_request,
        "spec": {
            "name": "http_request",
            "description": "Requete HTTP generique vers une API publique (GET/POST/PUT/PATCH/DELETE). Hotes prives interdits sauf admin.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string"},
                    "method": {"type": "string", "description": "GET par defaut"},
                    "headers": {"type": "object", "description": "En-tetes optionnels"},
                    "body": {"description": "Corps JSON ou texte pour POST/PUT/PATCH"},
                },
                "required": ["url"],
            },
        },
    },
    "memory_search": {
        "run": tool_memory_search,
        "spec": {
            "name": "memory_search",
            "description": "Recherche dans la memoire vectorielle des conversations passees de l'utilisateur.",
            "parameters": {
                "type": "object",
                "properties": {"query": {"type": "string"}},
                "required": ["query"],
            },
        },
    },
    "wikipedia": {
        "run": tool_wikipedia,
        "spec": {
            "name": "wikipedia",
            "description": "Cherche un article Wikipedia (fr) et retourne son resume.",
            "parameters": {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]},
        },
    },
    "weather": {
        "run": tool_weather,
        "spec": {
            "name": "weather",
            "description": "Meteo actuelle + previsions 3 jours pour une ville (Open-Meteo, sans cle).",
            "parameters": {"type": "object", "properties": {"city": {"type": "string"}}, "required": ["city"]},
        },
    },
    "rss_fetch": {
        "run": tool_rss_fetch,
        "spec": {
            "name": "rss_fetch",
            "description": "Lit un flux RSS/Atom et retourne les derniers articles (titre, lien, date).",
            "parameters": {
                "type": "object",
                "properties": {"url": {"type": "string"}, "max_items": {"type": "integer"}},
                "required": ["url"],
            },
        },
    },
    "youtube_transcript": {
        "run": tool_youtube_transcript,
        "spec": {
            "name": "youtube_transcript",
            "description": "Recupere la transcription d'une video YouTube (fr puis en).",
            "parameters": {"type": "object", "properties": {"url_or_id": {"type": "string"}}, "required": ["url_or_id"]},
        },
    },
    "generate_image": {
        "run": tool_generate_image,
        "spec": {
            "name": "generate_image",
            "description": "Genere une image (endpoint compatible OpenAI configure via IMAGE_API_BASE) et la sauve dans le workspace.",
            "parameters": {
                "type": "object",
                "properties": {"prompt": {"type": "string"}, "size": {"type": "string", "description": "ex: 1024x1024"}},
                "required": ["prompt"],
            },
        },
    },
    "deep_research": {
        "run": tool_deep_research,
        "spec": {
            "name": "deep_research",
            "description": "Recherche web approfondie multi-etapes sur un sujet -> rapport markdown structure avec sources. Plus long qu'une simple recherche.",
            "parameters": {"type": "object", "properties": {"topic": {"type": "string"}}, "required": ["topic"]},
        },
    },
    "python_exec": {
        "run": tool_python_exec,
        "spec": {
            "name": "python_exec",
            "description": "Execute du code Python dans le workspace (ADMIN seulement, timeout 30s, mode isole -I).",
            "parameters": {
                "type": "object",
                "properties": {"code": {"type": "string"}},
                "required": ["code"],
            },
        },
    },
    "code_execute": {
        "run": tool_code_execute,
        "spec": {
            "name": "code_execute",
            "description": "Execute du code Python dans un sandbox isole (disponible pour tous). Utile pour calculs, transformations de donnees, generation de texte programme. Timeout 10s. Acces reseau et systeme bloques.",
            "parameters": {
                "type": "object",
                "properties": {
                    "code": {"type": "string", "description": "Code Python a executer"},
                },
                "required": ["code"],
            },
        },
    },
    "shell": {
        "run": tool_shell,
        "spec": {
            "name": "shell",
            "description": "Execute une commande shell dans le workspace (ADMIN seulement, timeout 30s).",
            "parameters": {
                "type": "object",
                "properties": {"command": {"type": "string"}},
                "required": ["command"],
            },
        },
    },
    "note_add": {
        "run": tool_note_add,
        "spec": {
            "name": "note_add",
            "description": "Enregistre une note persistante pour l'utilisateur (memo, idee, information a retenir).",
            "parameters": {
                "type": "object",
                "properties": {"content": {"type": "string", "description": "Texte de la note"}},
                "required": ["content"],
            },
        },
    },
    "note_list": {
        "run": tool_note_list,
        "spec": {
            "name": "note_list",
            "description": "Liste les notes enregistrees de l'utilisateur (avec leur id).",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    "note_delete": {
        "run": tool_note_delete,
        "spec": {
            "name": "note_delete",
            "description": "Supprime une note par son id (prefixe accepte).",
            "parameters": {
                "type": "object",
                "properties": {"note_id": {"type": "string", "description": "Id (ou debut d'id) de la note"}},
                "required": ["note_id"],
            },
        },
    },
    "todo_add": {
        "run": tool_todo_add,
        "spec": {
            "name": "todo_add",
            "description": "Ajoute une tache a la todo-list persistante de l'utilisateur.",
            "parameters": {
                "type": "object",
                "properties": {"content": {"type": "string", "description": "Description de la tache"}},
                "required": ["content"],
            },
        },
    },
    "todo_list": {
        "run": tool_todo_list,
        "spec": {
            "name": "todo_list",
            "description": "Liste les taches de la todo-list (en cours par defaut).",
            "parameters": {
                "type": "object",
                "properties": {"include_done": {"type": "boolean", "description": "Inclure les taches terminees"}},
            },
        },
    },
    "todo_done": {
        "run": tool_todo_done,
        "spec": {
            "name": "todo_done",
            "description": "Marque une tache comme terminee par son id (prefixe accepte).",
            "parameters": {
                "type": "object",
                "properties": {"todo_id": {"type": "string", "description": "Id (ou debut d'id) de la tache"}},
                "required": ["todo_id"],
            },
        },
    },
    "email_send": {
        "run": tool_email_send,
        "spec": {
            "name": "email_send",
            "description": "Envoie un email (texte) via SMTP. Necessite SMTP_HOST/SMTP_USER/SMTP_PASSWORD en variables d'environnement.",
            "parameters": {
                "type": "object",
                "properties": {
                    "to": {"type": "string", "description": "Adresse du destinataire"},
                    "subject": {"type": "string"},
                    "body": {"type": "string", "description": "Corps du message (texte brut)"},
                },
                "required": ["to", "subject", "body"],
            },
        },
    },
    # --- v1.3 ----------------------------------------------------------------
    "memory_save": {
        "run": tool_memory_save,
        "spec": {
            "name": "memory_save",
            "description": "Memorise durablement un fait important dans la memoire vectorielle persistante.",
            "parameters": {
                "type": "object",
                "properties": {
                    "fact": {"type": "string", "description": "Le fait a memoriser (phrase courte et precise)"},
                },
                "required": ["fact"],
            },
        },
    },
    "session_search": {
        "run": tool_session_search,
        "spec": {
            "name": "session_search",
            "description": "Recherche plein texte (FTS5) dans toutes les conversations passees.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "Termes de recherche"},
                    "limit": {"type": "integer", "description": "Nombre max de resultats (defaut 8, max 20)"},
                },
                "required": ["query"],
            },
        },
    },
    "patch_file": {
        "run": tool_patch_file,
        "spec": {
            "name": "patch_file",
            "description": "Remplace un bloc de texte dans un fichier du workspace par un nouveau contenu.",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "Chemin relatif du fichier dans le workspace"},
                    "old_str": {"type": "string", "description": "Texte exact a remplacer (laisser vide pour ecraser tout le fichier)"},
                    "new_str": {"type": "string", "description": "Nouveau texte"},
                },
                "required": ["path", "new_str"],
            },
        },
    },
    "skill_save": {
        "run": tool_skill_save,
        "spec": {
            "name": "skill_save",
            "description": "Sauvegarde une procedure reutilisable (skill) en base. Permet a l'agent de retrouver ses methodes de travail.",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Identifiant court du skill (ex: redaction_email)"},
                    "description": {"type": "string", "description": "Courte description du skill"},
                    "body": {"type": "string", "description": "Corps de la procedure (texte libre, instructions pas-a-pas)"},
                },
                "required": ["name", "body"],
            },
        },
    },
    "skill_list": {
        "run": tool_skill_list,
        "spec": {
            "name": "skill_list",
            "description": "Liste tous les skills disponibles pour cet utilisateur.",
            "parameters": {"type": "object", "properties": {}},
        },
    },
    "skill_run": {
        "run": tool_skill_run,
        "spec": {
            "name": "skill_run",
            "description": "Rappelle le corps d'un skill sauvegarde pour l'executer.",
            "parameters": {
                "type": "object",
                "properties": {
                    "name": {"type": "string", "description": "Nom du skill a executer"},
                },
                "required": ["name"],
            },
        },
    },

    "image_generate": {
        "run": tool_image_generate,
        "spec": {
            "name": "image_generate",
            "description": "Genere une image a partir d'un prompt textuel via le provider OpenAI-compatible configure (DALL-E, etc.). Retourne l'URL de l'image.",
            "parameters": {
                "type": "object",
                "properties": {
                    "prompt": {"type": "string", "description": "Description precise de l'image a generer"},
                    "size": {"type": "string", "description": "Taille : 256x256 | 512x512 | 1024x1024 | 1792x1024 | 1024x1792", "default": "1024x1024"},
                    "quality": {"type": "string", "description": "standard ou hd (DALL-E 3 uniquement)", "default": "standard"},
                    "n": {"type": "integer", "description": "Nombre d'images (1-4)", "default": 1},
                },
                "required": ["prompt"],
            },
        },
    },
    "page_summary": {
        "run": tool_page_summary,
        "spec": {
            "name": "page_summary",
            "description": "Recupere et extrait le contenu textuel d'une page web a partir de son URL. Utile pour lire un article, une documentation ou repondre a une question sur un site.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {"type": "string", "description": "URL complete de la page (https://...)"},
                    "question": {"type": "string", "description": "Question optionnelle sur le contenu de la page"},
                },
                "required": ["url"],
            },
        },
    },
}


def openai_tool_specs(is_admin: bool, plan: str = "premium") -> list[dict]:
    """Specs au format OpenAI selon le plan."""
    specs = []
    for name, t in TOOLS.items():
        if not plans.tool_allowed(name, plan, is_admin):
            continue
        specs.append({"type": "function", "function": t["spec"]})
    return specs


async def execute_tool(name: str, args: dict, ctx: dict) -> str:
    tool = TOOLS.get(name)
    if not tool:
        return f"Outil inconnu : {name}"
    if name in plans.ADMIN_TOOLS and not bool(ctx.get("is_admin")):
        return "Administrateur requis pour utiliser cet outil."
    if not plans.tool_allowed(name, ctx.get("plan"), bool(ctx.get("is_admin"))):
        return "Premium requis pour utiliser les outils de l'agent."
    runner: Callable[[dict, dict], Awaitable[str]] = tool["run"]
    try:
        return await runner(args, ctx)
    except Exception as e:  # garde-fou : un outil ne doit jamais tuer la boucle
        return f"Erreur outil {name} : {e}"
